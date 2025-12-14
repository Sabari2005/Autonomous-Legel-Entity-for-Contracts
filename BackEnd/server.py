from fastapi import FastAPI, File, UploadFile, WebSocket, HTTPException, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
import html2text
import imaplib
import email
from email.header import decode_header
from datetime import datetime, timezone
import time
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from typing import List
from fastapi import Header
from typing import List, Tuple, Dict
from motor.motor_asyncio import AsyncIOMotorClient
from motor.motor_asyncio import AsyncIOMotorClient, AsyncIOMotorGridFSBucket
import mimetypes
from bson import ObjectId
import gridfs
import ast
from fastapi.responses import FileResponse
from fastapi.responses import HTMLResponse
import re
import torch
from transformers import BertTokenizer
from fastapi import Request
import shap
from groq import Groq
from fastapi.staticfiles import StaticFiles
import shutil
import subprocess
import os
import uuid
import pdfplumber
import multiprocessing
import fitz  
import docx
import torch
from transformers import BertTokenizer, BertForSequenceClassification, AdamW, get_linear_schedule_with_warmup
from sklearn.preprocessing import LabelEncoder
from torch.utils.data import DataLoader, TensorDataset, RandomSampler, SequentialSampler, random_split
import numpy as np
import pandas as pd
import re
from nltk.corpus import stopwords
import nltk
from pyngrok import ngrok
import random
import time
import datetime
import uvicorn
import shap
import groq
import asyncio
import socket
import httpx
from neo4j import GraphDatabase
from pyvis.network import Network
from sentence_transformers import SentenceTransformer, util
from PyPDF2 import PdfReader
from docx import Document
from langchain_ollama import OllamaLLM


import gc
import torch

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["POST", "GET", "OPTIONS"], 
    allow_headers=["*"],
)
CONTRACT_FILE_PATH="/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/contract.html"
RECAPTCHA_SECRET = "6LefvhUrAAAAAAr-ORKzTQ2vOOCKumuOLCTpo8o4"

@app.post("/verify-recaptcha")
async def verify_recaptcha(request: Request):
    data = await request.json()
    token = data.get("token")
    
    async with httpx.AsyncClient() as client:
        response = await client.post(
            "https://www.google.com/recaptcha/api/siteverify",
            data={
                "secret": RECAPTCHA_SECRET,
                "response": token
            }
        )
    result = response.json()
    return {"success": result.get("success", False)}


uri = "LINK TO CONNECT TO NEWO4J DATABASE"
username = "neo4j"
password = "PASSWORD OF NEO4J"
driver = GraphDatabase.driver(uri, auth=(username, password))
similarity_model = SentenceTransformer('all-MiniLM-L6-v2')


def clean_text(text):
    text = text.lower()
    text = re.sub(r"[^a-zA-Z?.!,¿]+", " ", text)
    text = re.sub(r"http\S+", "", text)
    punctuations = '@#!?+&*[]-%.:/();$=><|{}^' + "'`" + '_'
    for p in punctuations:
        text = text.replace(p, '')
    sw = stopwords.words('english')
    text = [word.lower() for word in text.split() if word.lower() not in sw]
    text = " ".join(text)
    return text


def predict(texts):
    if isinstance(texts, str):
        texts = [texts]
    elif isinstance(texts, np.ndarray):
        texts = list(texts)

    inputs = tokenizer(texts, return_tensors='pt', padding=True, truncation=True, max_length=512)
    input_ids = inputs['input_ids'].to(device)
    attention_mask = inputs['attention_mask'].to(device)

    with torch.no_grad():
        logits_level = model(input_ids, attention_mask=attention_mask)["logits_level"]
    
    return torch.softmax(logits_level, dim=-1).cpu().numpy()


def predict_risk(text):
    cleaned_text = clean_text(text)
    encoded_dict = tokenizer.encode_plus(
        cleaned_text, 
        add_special_tokens=True,
        max_length=512,
        padding='max_length',
        truncation=True,
        return_attention_mask=True,
        return_tensors='pt'
    )
    input_ids = encoded_dict['input_ids'].to(device)
    attention_mask = encoded_dict['attention_mask'].to(device)

    with torch.no_grad():
        outputs = model(input_ids, attention_mask=attention_mask)
  
    predicted_level = np.argmax(outputs["logits_level"].cpu().numpy(), axis=1).flatten()[0]
    predicted_category = np.argmax(outputs["logits_category"].cpu().numpy(), axis=1).flatten()[0]

    return (
        risk_level_mapping.get(predicted_level, "Unknown"),
        risk_category_mapping.get(predicted_category, "Unknown")
    )

def split_into_legal_chunks(text, max_length=512):

    paragraphs = re.split(r'\n', text.strip())
    chunks = []
    current_chunk = ""
    
    for para in paragraphs:
        temp_chunk = f"{current_chunk}\n\n{para}" if current_chunk else para
        tokens = tokenizer.tokenize(temp_chunk)
        
        if len(tokens) <= max_length:
            current_chunk = temp_chunk
        else:

            if current_chunk == "":
                sentences = re.split(r'(?<=[.!?])\s+', para)
                current_sentence_chunk = ""
                for sent in sentences:
                    temp_sent_chunk = f"{current_sentence_chunk} {sent}".strip()
                    sent_tokens = tokenizer.tokenize(temp_sent_chunk)
                    
                    if len(sent_tokens) <= max_length:
                        current_sentence_chunk = temp_sent_chunk
                    else:
                        chunks.append(current_sentence_chunk)
                        current_sentence_chunk = sent
                
                if current_sentence_chunk:
                    chunks.append(current_sentence_chunk)
            else:
                chunks.append(current_chunk)
                current_chunk = para
    
    if current_chunk:
        chunks.append(current_chunk)
    
    return chunks

def explain_chunk(chunk, risk_level, risk_category):
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": f"""
                You are a legal risk expert. Analyze the following contract text chunk and provide a detailed explanation of:
                1. Why this text was classified as {risk_level} risk
                2. The specific clauses or phrases that contributed to this risk level
                3. The implications of this risk category ({risk_category})
                
                Please provide the content using appropriate HTML tags instead of Markdown format.
#                 Ensure proper structure with paragraphs (`<p>`), bold (`<strong>`), lists (`<ul>`, `<li>`), and other necessary elements where applicable.
#                 Exclude explanations and provide only the final formatted HTML content.
                
                Text chunk:
                {chunk}

                Exclude explanation
                """
            }
        ],
        model="llama-3.3-70b-versatile",
    )
    return chat_completion.choices[0].message.content.strip()
def modifier_model(text, risk_level, risk_category):
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": f"""
                You are a legal risk mitigation expert. Modify the given contract text by removing or rewording only high or medium-risk clauses to make them balanced for both parties. Do not change the remaining text.  
                
                Risk Level: {risk_level}  
                Risk Category: {risk_category}  
                
                Text:  
                {text}  
                
                Provide without explanations, comments, or additional context.
                """
            }
        ],
        model="llama-3.3-70b-versatile",
    )
    return chat_completion.choices[0].message.content.strip()

def format_legal_text(modified_text):
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": f"""
                Format the given text using appropriate HTML tags while preserving its original content. Apply suitable tags for headings, paragraphs, lists, and emphasis where needed. Do not alter the wording or structure of the text.  
                
                Text:  
                {modified_text}  
                
                Provide the formatted text without explanations or additional comments.
                """
            }
        ],
        model="llama-3.3-70b-versatile",
    )
    return chat_completion.choices[0].message.content.strip()

def process_chunk(chunk):

    inputs = tokenizer(chunk, return_tensors="pt", truncation=True, max_length=512).to(device)

    risk_level, risk_category = predict_risk(chunk)
    class_names = label_encoder_level.classes_.tolist()

    explainer = shap.Explainer(predict, tokenizer,output_names=class_names)
    torch.backends.cudnn.benchmark = True 
    torch.set_num_threads(multiprocessing.cpu_count())  
    explanation = explain_chunk(chunk, risk_level, risk_category)
    try:
        if torch.cuda.is_available():
            with torch.amp.autocast(device_type='cuda', dtype=torch.float16): 
                shap_values = explainer([chunk])
        else:
            shap_values = explainer([chunk])
        shap_plot = shap.plots.text(shap_values, display=False)
    except Exception as e:
        print(f"SHAP analysis failed for chunk: {str(e)}")
        shap_values = None
        shap_plot = "<p>SHAP analysis unavailable</p>"
    if risk_level in ["High", "Medium","Low"]:
        modified_text = modifier_model(chunk, risk_level, risk_category)
        formatted_text = format_legal_text(modified_text)
    else:
        modified_text = chunk
        formatted_text = format_legal_text(chunk)

    return {
        "original": chunk,
        "risk_level": risk_level,
        "risk_category": risk_category,
        "shap_values": shap_values,
        "shap_plot": shap_plot,
        "explanation": explanation,
        "modified": modified_text,
        "formatted":formatted_text
    }

def fetch_all_clauses():
    with driver.session() as session:
        result = session.run("MATCH (c:Clause) RETURN c.text AS text, c.risk_level AS risk_level, c.risk_category AS risk_category")
        return [record.data() for record in result]

def find_similar_clauses(clause_text, risk_level, risk_category, top_k=5):
    input_embedding = similarity_model.encode(clause_text, convert_to_tensor=True)
    clauses = fetch_all_clauses()

    risk_level = risk_level.lower().strip()
    risk_category = risk_category.lower().strip()

    filtered_clauses = [
        c for c in clauses
        if c["risk_level"].lower().strip() == risk_level and
           c["risk_category"].lower().strip() == risk_category
    ]

    if not filtered_clauses:
        return [], clause_text

    clause_texts = [clause['text'] for clause in filtered_clauses]
    clause_embeddings = similarity_model.encode(clause_texts, convert_to_tensor=True)

    scores = util.pytorch_cos_sim(input_embedding, clause_embeddings)[0]
    top_results = scores.topk(k=min(top_k, len(filtered_clauses)))

    matched = []
    for score, idx in zip(top_results.values, top_results.indices):
        clause = filtered_clauses[idx]
        clause["similarity"] = float(score)
        matched.append(clause)

    return matched, clause_text


def process_text(input_text):
    chunks = split_into_legal_chunks(input_text)
    chunks_data = []
    
    print(f"\nProcessing {len(chunks)} text chunks...")
    
    for i, chunk in enumerate(chunks):
        print(f"\nProcessing chunk {i+1}/{len(chunks)}")
        print(f"Chunk length: {len(chunk)} characters")
        
        risk_level, risk_category = predict_risk(chunk)
        print(f"Predicted Risk: {risk_level}")
        print(f"Predicted Category: {risk_category}")
        
        similar_clauses, _ = find_similar_clauses(chunk, risk_level, risk_category)
        
        if similar_clauses:
            print(f"Found {len(similar_clauses)} similar clauses")
            chunks_data.append((chunk, risk_level, risk_category, similar_clauses))
        else:
            print("No similar clauses found")
            chunks_data.append((chunk, risk_level, risk_category, []))
    
    return chunks_data


def save_results_as_html(results, filename="/teamspace/studios/this_studio/Uvarajan/Whole_Server/legal_risk_analysis_report.html"):
    html_content = """
    <html>
    <head>
        <style>
            body { font-family: Arial, sans-serif }
            .chunk-container { border: 1px solid #ddd; margin-bottom: 20px; padding: 15px; }
            .risk-high { color: #e74c3c; font-weight: bold; }
            .risk-medium { color: #f39c12; font-weight: bold; }
            .risk-low { color: #2ecc71; font-weight: bold; }
            .explanation { background-color: #f8f9fa; padding: 10px; margin-top: 10px;  text-align: justify;
  text-align-last: justify; }
            .modified_web { background-color: #e8f5e9; padding: 10px; margin-top: 10px;  text-align: justify;
  text-align-last: justify; }
            .formatted { background-color: #e3f2fd; padding: 10px; margin-top: 10px; 
                         white-space: pre-wrap; font-family: 'Courier New', monospace; }
            .shap-plot { margin: 15px 0;  text-align: justify;
  text-align-last: justify; }
            h2 { color: #2c3e50; border-bottom: 1px solid #eee; padding-bottom: 5px; }
            .original-text { white-space: pre-wrap; background-color: #f5f5f5; padding: 10px;  text-align: justify;
  text-align-last: justify; }
            .error { color: #e74c3c; background-color: #fde8e8; padding: 10px; }
        </style>
    </head>
    <body>
    """
    
    full_doc_original = []
    full_doc_modified = []
    full_doc_formatted = []
    
    for i, result in enumerate(results, 1):
        risk_class = f"risk-{result['risk_level'].lower()}"
        
        html_content += f"""
        <div class="chunk-container">
        <div style="display:flex;justify-content:space-between;align-items:center;">
                <p><strong>Risk Level:</strong> <span class="{risk_class}">{result['risk_level']}</span></p>
                <p><strong>Risk Category:</strong> {result['risk_category']}</p>
            </div>
            
            <div class="shap-plot">{result['shap_plot']}</div>
            
            <h3>Risk Explanation</h3>
            <div class="explanation">{result['explanation']}</div>
            

        </div>
        """
        
        full_doc_original.append(result['original'])
        full_doc_modified.append(result['modified'])
        full_doc_formatted.append(result['formatted'])
        html_content += """
        
        </body>
        </html>
        """
        with open(filename, "w", encoding="utf-8") as f:
            f.write(html_content)
        print(f"Report saved to {filename}")

    return full_doc_original, full_doc_modified, full_doc_formatted, html_content


def extract_text_from_file(file_path: str, file_ext: str) -> str:
    extracted_text = ""
    if file_ext == "pdf":
        with open(file_path, "rb") as pdf_file:
            reader = PdfReader(pdf_file)
            extracted_text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif file_ext == "docx":
        doc = Document(file_path)
        extracted_text = "\n".join([para.text for para in doc.paragraphs])
    return extracted_text.strip()
def docx_to_pdf(input_path: str, output_pdf: str):
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", input_path, "--outdir", os.path.dirname(output_pdf)],
            check=True
        )
        print(f"Converted '{input_path}' to PDF: {output_pdf}")
    except subprocess.CalledProcessError as e:
        print(f"Error converting DOCX to PDF: {e}")
def pdf_to_html(pdf_path: str, output_html: str) -> str:
    html_content = "<html><body>"
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                html_content += "<p>" + text.replace("\n", "<br>") + "</p>"
    html_content += "</body></html>"

    with open(output_html, "w", encoding="utf-8") as f:
        f.write(html_content)
    
    print(f"✅ Converted '{pdf_path}' to HTML: {output_html}")
    return html_content  
def clean_extracted_text(text):
    text = text.replace('\r', '')
    text = re.sub(r'\n+', '\n', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text

import os
import easyocr
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from PyPDF2 import PdfReader
from docx import Document
from bs4 import BeautifulSoup

class TextExtractor:
    def __init__(self):
        self.reader = easyocr.Reader(['en'], gpu=self._check_gpu())
    
    def _check_gpu(self):
        try:
            import torch
            return torch.cuda.is_available()
        except:
            return False

    def _extract_from_image(self, image_path):
        try:
            result = self.reader.readtext(image_path, detail=0)
            return " ".join(result) if result else pytesseract.image_to_string(Image.open(image_path))
        except Exception as e:
            raise Exception(f"Image extraction failed: {e}")

    def _extract_from_pdf(self, pdf_path):
        text = ""
        try:
            with open(pdf_path, 'rb') as f:
                pdf = PdfReader(f)
                text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
            
            if not text.strip():
                images = convert_from_path(pdf_path)
                for i, img in enumerate(images):
                    img_path = f"temp_pdf_page_{i}.jpg"
                    img.save(img_path, 'JPEG')
                    text += self._extract_from_image(img_path) + "\n"
                    os.remove(img_path)
            return text
        except Exception as e:
            raise Exception(f"PDF extraction failed: {e}")

    def _extract_from_docx(self, docx_path):
        try:
            doc = Document(docx_path)
            return "\n".join(para.text for para in doc.paragraphs if para.text)
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {e}")

    def _extract_from_html(self, html_path):
        try:
            with open(html_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f, 'html.parser')
                return soup.get_text(separator='\n', strip=True)
        except Exception as e:
            raise Exception(f"HTML extraction failed: {e}")

    def extract(self, file_path):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = file_path.lower().split('.')[-1]
        
        if ext in ['jpg', 'jpeg', 'png', 'bmp']:
            return self._extract_from_image(file_path)
        elif ext == 'pdf':
            return self._extract_from_pdf(file_path)
        elif ext == 'docx':
            return self._extract_from_docx(file_path)
        elif ext == 'html':
            return self._extract_from_html(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

import traceback
def process_document(file_location: str, file_ext: str):
    def create_combined_graph(chunks_data, output_file="/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/kg/combined_knowledge_graph.html"):
        net = Network(height="700px", width="100%", directed=True, notebook=True)
        net.force_atlas_2based()
        
        for chunk_idx, (chunk_text, risk_level, risk_category, similar_clauses) in enumerate(chunks_data):
            chunk_node_id = f"Chunk_{chunk_idx + 1}"
            risk_color = 'red' if risk_level.lower() == 'high' else 'orange' if risk_level.lower() == 'medium' else 'green'
            net.add_node(chunk_node_id, 
                        label=f"Chunk {chunk_idx + 1}",
                        color="lightblue",
                        title=f"{risk_level} risk - {risk_category}\n\n{chunk_text[:200]}...")
            risk_node = f"{chunk_node_id}_Risk_Level"
            category_node = f"{chunk_node_id}_Category"
            
            net.add_node(risk_node, 
                        label=f"Risk: {risk_level}",
                        color=risk_color,
                        title=f"Risk Level: {risk_level}")
            
            net.add_node(category_node,
                        label=f"Category: {risk_category}",
                        color="purple",
                        title=f"Category: {risk_category}")
            net.add_edge(chunk_node_id, risk_node, label="HAS_RISK")
            net.add_edge(chunk_node_id, category_node, label="HAS_CATEGORY")
            for i, clause in enumerate(similar_clauses[:3]):
                clause_id = f"{chunk_node_id}_Match_{i+1}"
                net.add_node(clause_id,
                            label=f"Match {i+1}",
                            color="lightcoral",
                            title=f"Similarity: {clause['similarity']:.2f}\n\n{clause['text'][:200]}...")
                net.add_edge(chunk_node_id, clause_id, label="SIMILAR_TO")
        
        net.show_buttons(filter_=["physics"])
        net.save_graph(output_file)
        print(f"\nCombined knowledge graph saved to {output_file}")
        return output_file
    extractor = TextExtractor()
    
    extracted_text = extractor.extract(file_location)  
    print(extracted_text)
    cleaned_text = clean_extracted_text(extracted_text)
    chunks = split_into_legal_chunks(cleaned_text)

    results = []
    chunks_data = process_text(cleaned_text)
    kg_out=create_combined_graph(chunks_data)
    for chunk in chunks:
        try:
            results.append(process_chunk(chunk))
        except Exception as e:
            print(f"Error processing chunk: {str(e)}")
            traceback.print_exc()
            results.append({
                "original": chunk,
                "risk_level": "Error",
                "risk_category": "Processing Error",
                "shap_values": None,
                "explanation": f"Analysis failed: {str(e)}",
                "modified": chunk
                
            })
    full_doc_original, full_doc_modified, full_doc_formatted, html_content = save_results_as_html(results)

    html_path = ""
    explained_text = ""
    if file_ext == "docx":
        pdf_path = file_location.replace(".docx", ".pdf")
        docx_to_pdf(file_location, pdf_path)
        html_path = pdf_path.replace(".pdf", ".html")
        explained_text = pdf_to_html(pdf_path, html_path)
        os.remove(pdf_path) 

    elif file_ext == "pdf":
        html_path = file_location.replace(".pdf", ".html")
        explained_text = pdf_to_html(file_location, html_path)

    return {
        "extracted_text": cleaned_text,
        "modified_text": full_doc_formatted,
        "explained_text": html_content,
        "htmlUrl": html_path,
        "kg_out": kg_out
    }



UPLOAD_DIR = "/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
app.mount("/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")


processed_files = {}
active_sessions = {}

nltk.download('stopwords')
groq_API_KEY  = "GROQ API KEY"
client =Groq(api_key=groq_API_KEY)

df = pd.read_csv("/teamspace/studios/this_studio/Uvarajan/Whole_Server/Final_Full_Dataset.csv")

df = df.dropna(subset=["Risk Level"])

label_encoder_level = LabelEncoder()
label_encoder_category = LabelEncoder()

df['Risk Level'] = label_encoder_level.fit_transform(df['Risk Level'])
df['Risk Category'] = label_encoder_category.fit_transform(df['Risk Category'])


risk_level_mapping = {i: label for i, label in enumerate(label_encoder_level.classes_)}
risk_category_mapping = {i: label for i, label in enumerate(label_encoder_category.classes_)}

print("Risk Level Encodings:")
for encoded, original in risk_level_mapping.items():
    print(f"{original}: {encoded}")

print("\nRisk Category Encodings:")
for encoded, original in risk_category_mapping.items():
    print(f"{original}: {encoded}")


tokenizer = BertTokenizer.from_pretrained('bert-large-uncased', do_lower_case=True, model_max_length=512)

class BertForMultiTaskClassification(BertForSequenceClassification):
    def __init__(self, config, num_labels_level, num_labels_category):
        super().__init__(config)
        self.num_labels_level = num_labels_level
        self.num_labels_category = num_labels_category
        self.classifier_level = torch.nn.Linear(config.hidden_size, num_labels_level)
        self.classifier_category = torch.nn.Linear(config.hidden_size, num_labels_category)

    def forward(self, input_ids, attention_mask=None, labels_level=None, labels_category=None):
        outputs = self.bert(input_ids, attention_mask=attention_mask)
        pooled_output = outputs[1]

        logits_level = self.classifier_level(pooled_output)
        logits_category = self.classifier_category(pooled_output)

        loss = None
        if labels_level is not None and labels_category is not None:
            loss_fct = torch.nn.CrossEntropyLoss()
            loss_level = loss_fct(logits_level.view(-1, self.num_labels_level), labels_level.view(-1))
            loss_category = loss_fct(logits_category.view(-1, self.num_labels_category), labels_category.view(-1))
            loss = loss_level + loss_category

        return {
            "logits_level": logits_level,
            "logits_category": logits_category,
            "loss": loss
        }


class RiskAnalyzer:
    def __init__(self, model_path: str = '/teamspace/studios/this_studio/Uvarajan/Whole_Server/Risk_analysis_BERT.pth'):
        try:
            self.device = torch.device("cpu")
            df = pd.read_csv("/teamspace/studios/this_studio/Uvarajan/Whole_Server/Final_Full_Dataset.csv")
            df = df.dropna(subset=["Risk Level"])
            
            self.label_encoder_level = LabelEncoder()
            self.label_encoder_category = LabelEncoder()
            df['Risk Level'] = self.label_encoder_level.fit_transform(df['Risk Level'])
            df['Risk Category'] = self.label_encoder_category.fit_transform(df['Risk Category'])
            
            self.risk_level_mapping = {i: label for i, label in enumerate(self.label_encoder_level.classes_)}
            self.risk_category_mapping = {i: label for i, label in enumerate(self.label_encoder_category.classes_)}
            
            self.tokenizer = BertTokenizer.from_pretrained('bert-large-uncased', do_lower_case=True, model_max_length=512)
            self.model = self._load_model(model_path)
        except Exception as e:
            raise RuntimeError(f"Failed to initialize RiskAnalyzer: {str(e)}")

    def _load_model(self, model_path: str) -> BertForSequenceClassification:
        try:
            model = BertForMultiTaskClassification.from_pretrained(
                "bert-large-uncased",
                num_labels_level=len(self.label_encoder_level.classes_),
                num_labels_category=len(self.label_encoder_category.classes_))
            model.load_state_dict(torch.load(model_path, map_location=torch.device('cpu')))
            model.to(self.device)
            model.eval()
            return model
        except Exception as e:
            raise RuntimeError(f"Failed to load model: {str(e)}")

    def clean_text(self, text: str) -> str:
        try:
            text = text.lower()
            text = re.sub(r"[^a-zA-Z?.!,¿]+", " ", text)
            text = re.sub(r"http\S+", "", text)
            punctuations = '@#!?+&*[]-%.:/();$=><|{}^' + "'`" + '_'
            for p in punctuations:
                text = text.replace(p, '')
            sw = stopwords.words('english')
            text = [word.lower() for word in text.split() if word.lower() not in sw]
            return " ".join(text)
        except Exception as e:
            raise RuntimeError(f"Text cleaning failed: {str(e)}")

    def analyze_risk(self, text: str) -> Tuple[str, str]:
        try:
            cleaned_text = self.clean_text(text)
            encoded_dict = self.tokenizer.encode_plus(
                cleaned_text,
                add_special_tokens=True,
                max_length=512,
                padding='max_length',
                truncation=True,
                return_attention_mask=True,
                return_tensors='pt'
            )
            
            input_ids = encoded_dict['input_ids'].to(self.device)
            attention_mask = encoded_dict['attention_mask'].to(self.device)
            
            with torch.no_grad():
                outputs = self.model(input_ids, attention_mask=attention_mask)
            
            predicted_level = np.argmax(outputs["logits_level"].cpu().numpy(), axis=1).flatten()[0]
            predicted_category = np.argmax(outputs["logits_category"].cpu().numpy(), axis=1).flatten()[0]
            
            return (
                str(self.risk_level_mapping.get(predicted_level, "Unknown")),
                str(self.risk_category_mapping.get(predicted_category, "Unknown"))
            )
        except Exception as e:
            raise RuntimeError(f"Risk analysis failed: {str(e)}")

    def unload_model(self):
        try:
            if hasattr(self, 'model'):
                self.model.to('cpu')
                del self.model
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                if hasattr(self, 'tokenizer'):
                    del self.tokenizer
                    
                return True
            return False
        except Exception as e:
            print(f"Error unloading model: {str(e)}")
            return False

class ContinuousContractSummarizer:
    def __init__(self, model: str = "llama3.3"):
        try:
            self.llm = OllamaLLM(model=model)
            self.risk_analyzer = RiskAnalyzer()
            print(f"Risk Analyzer device: CPU")
            if torch.cuda.is_available():
                print(f"LLM will use GPU: {torch.cuda.get_device_name(0)}")
            else:
                print("LLM will use CPU")
        except Exception as e:
            raise RuntimeError(f"Failed to initialize summarizer: {str(e)}")

    def _split_contract(self, text: str) -> List[str]:
        try:
            sections = re.split(r'\n\s*(ARTICLE|SECTION)\s+[IVXLCDM]+\s*\n', text)
            if len(sections) > 1:
                return sections
            return re.split(r'\n\d+\.\s+[A-Z][A-Z\s]+\n', text) or [text]
        except Exception as e:
            raise RuntimeError(f"Contract splitting failed: {str(e)}")

    def _generate_prompt(self, text: str, risk_info: Dict[str, str] = None, is_final: bool = False) -> str:
        risk_context = ""
        if risk_info:
            risk_context = f"""
            RISK ANALYSIS RESULTS:
            - Risk Level: {risk_info['level']}
            - Risk Category: {risk_info['category']}
            
            IMPORTANT: Pay special attention to high-risk sections and highlight any terms related to the identified risk category.
            """
        
        if is_final:
            return f"""
            {risk_context}
            SUMMARIZE THIS CONTRACT SECTION CLEARLY AND CONCISELY:
            RULES:
            - Keep the summary in *natural language* (no bullet points).
            - Highlight IMPORTANT INFO using *bold* (e.g., *Party A, **$10,000, **terminate, **30 days*).
            - Pay special attention to risk-related terms based on the risk analysis above.
            - Focus on:
            • Who is involved  
            • What must be done  
            • When deadlines/payments happen  
            • How to exit/terminate  
            • Any identified risks
            - Include only essential clauses (financial, legal obligations, dates, risks, termination).
            - Remove boilerplate, generic text, and formalities.
            INPUT:
            {text}
            Output: A short, readable paragraph summary with bold keywords.
            IF THE TEXT WAS NOT A CONTRACT OR AGREEMENT, RETURN:"I DO NOT UNDERSTAND THIS TEXT. IT IS NOT A CONTRACT OR AGREEMENT."
            """
        else:
            return f"""
            {risk_context}
            EXTRACT KEY LEGAL ELEMENTS FROM THE FOLLOWING CONTRACT TEXT:
            RULES:
            - Identify and capture:
            • *Involved parties*  
            • *Payment details* (amounts, due dates)  
            • *Termination conditions*  
            • *Key obligations* and *responsibilities*  
            • *Risks or conditions* (especially those matching the identified risk category)
            - Keep monetary values, deadlines, durations, and clause names *exactly as-is*
            - Highlight the intent without including generic legal boilerplate
            - Return a CLEAN, minimal version of the contract's core logic
            TEXT TO ANALYZE:
            {text}
            Output: Clean, reduced terms ready for summary.
            IF THE TEXT WAS NOT A CONTRACT OR AGREEMENT, RETURN:"I DO NOT UNDERSTAND THIS TEXT. IT IS NOT A CONTRACT OR AGREEMENT."
            """




    
    
    
    async def summarize(self, contract_text: str, session_id: str):
        try:
            if session_id not in active_sessions:
                return  

            chunks = self._split_contract(contract_text)
            extracted_terms = []

            for chunk in chunks:
                if session_id not in active_sessions:
                    return  

                try:
                    risk_level, risk_category = self.risk_analyzer.analyze_risk(chunk)
                    print(f"\n[Risk Analysis] Level: {risk_level}, Category: {risk_category}")
                    risk_info = {'level': risk_level, 'category': risk_category}
                    prompt = self._generate_prompt(chunk, risk_info)

                    print("\n[Processing Chunk]")
                    chunk_extract = ""
                    for chunk_output in self.llm.stream(prompt):  
                        chunk_extract += chunk_output

                    extracted_terms.append({
                        'text': chunk_extract,
                        'risk_level': risk_level,
                        'risk_category': risk_category
                    })

                except Exception as e:
                    print(f"\nError processing chunk: {str(e)}")
                    continue

            combined_terms = "\n".join(
                [f"[Risk: {item['risk_level']} - {item['risk_category']}]\n{item['text']}" for item in extracted_terms]
            )

            overall_risk_level, overall_risk_category = self.risk_analyzer.analyze_risk(contract_text)
            final_risk_info = {
                'level': overall_risk_level,
                'category': overall_risk_category
            }

            yield "data: \n\n================ FINAL SUMMARY ==================\n\n"
            print("\n\n================ FINAL SUMMARY ==================")
            print(f"[Overall Risk Assessment] Level: {overall_risk_level}, Category: {overall_risk_category}\n")

            final_prompt = self._generate_prompt(combined_terms, final_risk_info, is_final=True)

            async for chunk in self.llm.astream(final_prompt):  
                if session_id not in active_sessions:
                    return
                yield f"data: {chunk}\n\n"
                await asyncio.sleep(0.01)

            yield "data:#3jk*24 \n\n"  #

        except Exception as e:
            raise RuntimeError(f"Summarization failed: {str(e)}")




model = None
tokenizer = None
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
current_model = None  


def load_model(model_type):
    global current_model, model, tokenizer, device

    if current_model == model_type:
        return  
    if current_model in ["bert", "llama"]:
        del model
        del tokenizer
        gc.collect()
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
    if model_type == "bert":
        tokenizer = BertTokenizer.from_pretrained('bert-large-uncased', do_lower_case=True)
        model = BertForMultiTaskClassification.from_pretrained(
            "bert-large-uncased",
            num_labels_level=len(label_encoder_level.classes_),
            num_labels_category=len(label_encoder_category.classes_)
        )
        model.load_state_dict(torch.load(
            '/teamspace/studios/this_studio/Uvarajan/Whole_Server/Risk_analysis_BERT.pth',
            map_location=device
        ))
        model = model.to(device)
        model.eval()
        current_model = "bert"
        print("✅ BERT model loaded successfully.")

    elif model_type == "llama":
        risk=RiskAnalyzer()
        risk.unload_model()

@app.on_event("startup")
async def startup_db():
    app.state.mongo_client = AsyncIOMotorClient(
        "MONGO DB URI"  #store the templates in the mongodb to access it faster and aywhere
    )
    app.state.db3 = app.state.mongo_client["filesDB"]
    app.state.fs_files_collection = app.state.db3["fs.files"]

@app.on_event("shutdown")
async def shutdown_db():
    app.state.mongo_client.close()


def get_user_db(user_id: str):
    db_name = f"filesDB_{user_id}"  
    db = app.state.mongo_client[db_name]
    app.state.db3 = db
    app.state.fs_files_collection = db["fs.files"]
    return db



async def upload_file_to_db(file_path: str, user_id: str):
    get_user_db(user_id)  
    filename = os.path.basename(file_path)
    file_ext = os.path.splitext(filename)[1][1:]
    file_size = os.path.getsize(file_path)
    now = datetime.datetime.utcnow()

    with open(file_path, "rb") as f:
        file_data = f.read()

    fs_bucket = AsyncIOMotorGridFSBucket(app.state.db3)

    file_id = await fs_bucket.upload_from_stream(
        filename,
        file_data,
        metadata={
            "user_id": user_id,
            "type": file_ext,
            "size": file_size,
            "lastModified": now
        }
    )


    print(f"✅ File '{filename}' uploaded with ID: {file_id}")
    return str(file_id)


@app.get("/api/files")
async def get_files(userid: str = Header(...)): 
    get_user_db(userid)  
    print(userid)
    cursor = app.state.fs_files_collection.find(
    {"metadata.user_id": userid}, 
    sort=[("uploadDate", -1)]
)

    files = await cursor.to_list(length=None)
    file_list = []

    for file in files:
        file_list.append({
            "name": file.get("filename"),
            "size": file.get("metadata", {}).get("size"),
            "type": file.get("metadata", {}).get("type"),
            "lastModified": file.get("metadata", {}).get("lastModified").isoformat() 
                            if file.get("metadata", {}).get("lastModified") else None,
            "id": str(file.get("_id")),
            "url": f"/api/files/{file.get('_id')}"
        })

    return file_list


@app.get("/api/files/{file_id}")
async def download_file(file_id: str, userid: str = Header(...)):
    try:

        get_user_db(userid)

        file_obj = await app.state.fs_files_collection.find_one({"_id": ObjectId(file_id)})
        if not file_obj:
            raise HTTPException(status_code=404, detail="File not found")

        filename = file_obj.get("filename", "downloaded_file")
        content_type = file_obj.get("metadata", {}).get("type") or \
                       mimetypes.guess_type(filename)[0] or "application/octet-stream"

        fs_bucket = AsyncIOMotorGridFSBucket(app.state.db3)
        stream = await fs_bucket.open_download_stream(ObjectId(file_id))

        async def file_iterator():
            while True:
                chunk = await stream.read(4096)
                if not chunk:
                    break
                yield chunk

        headers = {
            "Content-Disposition": f'attachment; filename="{filename}"'
        }

        return StreamingResponse(file_iterator(), media_type=content_type, headers=headers)

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error downloading file: {str(e)}")

@app.get("/summarizer_text_sse")
async def summarizer_text_sse(text: str):

    headers = {
        "Content-Type": "text/event-stream",
        "Cache-Control": "no-cache",
        "Connection": "keep-alive",
    }

    session_id = str(uuid.uuid4())
    active_sessions[session_id] = True

    summarizer = ContinuousContractSummarizer()
    return StreamingResponse(summarizer.summarize(text, session_id), media_type="text/event-stream", headers={
        "Access-Control-Allow-Origin": "*",  
        "Access-Control-Allow-Credentials": "true"
    })



@app.post("/summerizer_upload_file")
async def summerizer_upload_file(request: Request,file: UploadFile = File(...)):
    file_ext = file.filename.split(".")[-1].lower()

    file_location = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_location, "wb") as buffer:
        buffer.write(await file.read())
    
    user_id = request.headers.get("userid")
    file_id = await upload_file_to_db(file_location,user_id)
    print("Uploaded file ID:", file_id)
    extractor = TextExtractor()
    try:
        extracted_text = extractor.extract(file_location)  
        print(extracted_text)
    except Exception as e:
        print(f"Error: {e}")
    file_id = str(uuid.uuid4())

    processed_files[file_id] = extracted_text

    return {"response": "File uploaded successfully", "file_id": file_id}


@app.get("/summarizer_file_sse")
async def summarizer_file_sse(file_id: str):

    headers = {
        "Content-Type": "text/event-stream",
        "Cache-Control": "no-cache",
        "Connection": "keep-alive",
    }

    if file_id not in processed_files:
        return StreamingResponse(iter(["event: error\ndata: Invalid file_id\n\n"]), headers=headers)

    contract_text = processed_files[file_id]
    session_id = file_id  
    active_sessions[session_id] = True

    summarizer = ContinuousContractSummarizer()
    return StreamingResponse(summarizer.summarize(contract_text, session_id), media_type="text/event-stream", headers={
            "Access-Control-Allow-Origin": "*", 
            "Access-Control-Allow-Credentials": "true"
        })


@app.post("/cancel_sse")
async def cancel_sse(session_id: str):

    if session_id in active_sessions:
        del active_sessions[session_id]
        return {"message": "SSE connection closed"}
    return {"message": "No active session"}, 400



@app.post("/upload/")
async def upload_file(file: UploadFile = File(...), userid: str = Header(...)):
    risk=RiskAnalyzer()
    risk.unload_model()
    load_model("bert")
    file_ext = file.filename.split(".")[-1].lower()


    file_location = os.path.join(UPLOAD_DIR, file.filename)
    
    with open(file_location, "wb") as buffer:
        buffer.write(await file.read())
    file_id = await upload_file_to_db(file_location,userid)
    print("Uploaded file ID:", file_id)
    doc_data = process_document(file_location, file_ext)

    return JSONResponse(content={
        "filename": file.filename,
        "fileUrl": file_location,
        "message": "File uploaded successfully",
        "extracted_text": doc_data["extracted_text"],
        "htmlUrl": doc_data["htmlUrl"],
        "original_text": doc_data["extracted_text"],
        "modified_text": doc_data["modified_text"],
        "explained_text": doc_data["explained_text"],
        "kg_out": doc_data["kg_out"]
    })

@app.post("/convert")
async def html_to_pdf(request: Request):

    try:
        data = await request.json()  
        html_content = data.get("html_content", "")

        if not html_content:
            raise HTTPException(status_code=400, detail="No HTML content provided")
        user_id = request.headers.get("userid")
        html_filename = f"/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/{uuid.uuid4()}.html"
        pdf_filename = f"/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/{uuid.uuid4()}.pdf"

        with open(html_filename, "w", encoding="utf-8") as f:
            f.write(html_content)
        print(html_content)

        subprocess.run(["wkhtmltopdf", html_filename, pdf_filename], check=True)
        file_id = await upload_file_to_db(pdf_filename,user_id)
        print("Uploaded file ID:", file_id)
        return FileResponse(pdf_filename, filename="output.pdf", media_type="application/pdf")

    except subprocess.CalledProcessError as e:
        raise HTTPException(status_code=500, detail=f"Conversion error: {str(e)}")

import markdown
from pathlib import Path
import json

import re


def generate_report_html(report_path):
    report_path = report_path

    with open(report_path, "r", encoding="utf-8") as f:
        md_content = f.read()

    html_body = markdown.markdown(md_content, extensions=["fenced_code", "tables"])

    html_body = re.sub(r'<a href="(.*?)">', r'<a href="\1" target="_blank">', html_body)

    final_html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>Company Unified Risk Report</title>
    <style>
        body {{ font-family: Arial, sans-serif; padding: 2rem; line-height: 1.6; }}
        h1, h2, h3 {{ color: #2c3e50; }}
        img {{ max-width: 100%; height: auto; margin: 1rem 0; }}
        pre {{ background: #f4f4f4; padding: 1rem; }}
        table {{
    border-collapse: collapse;
    width: 100%;
    margin-bottom: 1rem;
}}
th, td {{
    border: 1px solid black;
    padding: 8px;
    text-align: left;
}}
th {{
    background-color: #f2f2f2;
}}

    </style>
</head>
<body>
    {html_body}
    <div style="height: 60px;"></div>
    <script>
        window.onload = function () {{
            setTimeout(function() {{
                const height = document.body.scrollHeight;
                window.parent.postMessage({{ iframeHeight: height }}, '*');
            }}, 100);  // slight delay to ensure content is rendered
        }};
    </script>
</body>
</html>
"""

    output_path = Path("/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/html_report/Infosys_Unified_Risk_Report.html")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(final_html)
    return str(output_path)

corporation=""

import requests
import json
import pandas as pd
import numpy as np
from datetime import datetime, timedelta
from pathlib import Path
import logging
import re
import time
import matplotlib.pyplot as plt
from typing import Dict, List, Optional, Tuple
import subprocess
import os
from PyPDF2 import PdfMerger
from jinja2 import Template
from datetime import date
import asyncio
class FinancialRiskEvaluator:
    def __init__(self, finance_api_key: str, news_api_key: str):
        self.finance_api = finance_api_key
        self.news_api = news_api_key
        self.reports_folder = Path("/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Stock_analyses_output")
        self.reports_folder.mkdir(exist_ok=True)
        self.assessment_parameters = {
            'liquidity_ratio': {'caution': 1.5, 'critical': 1.0},
            'leverage_ratio': {'caution': 1.0, 'critical': 2.0},
            'profitability': {'caution': 0.05, 'critical': 0.0},
            'financial_health_score': {'caution': 3, 'critical': 1}
        }

        self.fraud_detection_weights = {
            'receivables_index': 0.5,
            'margin_index': 1.0,
            'asset_quality': 0.5,
            'growth_index': 1.0,
            'depreciation_rate': 0.5,
            'expense_index': 1.0,
            'leverage_change': 1.0,
            'accruals_ratio': 0.5
        }

    def retrieve_financial_statistics(self, ticker: str) -> dict:
        endpoint = f"https://www.alphavantage.co/query?function=OVERVIEW&symbol={ticker}&apikey={self.finance_api}"
        try:
            api_response = requests.get(endpoint, timeout=10)
            statistics = api_response.json()
            return {
                'valuation_ratio': float(statistics.get('PERatio', 0)),
                'net_margin': float(statistics.get('ProfitMargin', 0)),
                'liquidity_ratio': float(statistics.get('CurrentRatio', 0)),
                'leverage_ratio': float(statistics.get('DebtToEquity', 0)),
                'asset_returns': float(statistics.get('ReturnOnAssetsTTM', 0)),
                'equity_returns': float(statistics.get('ReturnOnEquityTTM', 0)),
                'company_valuation': float(statistics.get('MarketCapitalization', 0)),
                'annual_peak': float(statistics.get('52WeekHigh', 0)),
                'annual_trough': float(statistics.get('52WeekLow', 0))
            }
        except Exception as error:
            logging.error("Failed to retrieve financial data: %s", error)
            return {}

    def obtain_price_history(self, ticker: str) -> dict:

        api_url = f"https://www.alphavantage.co/query?function=TIME_SERIES_DAILY&symbol={ticker}&apikey={self.finance_api}&outputsize=full"
        try:
            price_response = requests.get(api_url, timeout=10)
            market_data = price_response.json()
            daily_prices = market_data.get('Time Series (Daily)', {})
            
            price_frame = pd.DataFrame.from_dict(daily_prices, orient='index')
            price_frame = price_frame.apply(pd.to_numeric)
            price_frame.index = pd.to_datetime(price_frame.index)
            price_frame = price_frame.sort_index()
            
            closing_prices = price_frame['4. close']
            recent_close = closing_prices.iloc[-1]
            medium_term_avg = closing_prices.rolling(50).mean()[-1]
            long_term_avg = closing_prices.rolling(200).mean()[-1]
            price_variability = closing_prices.pct_change().std() * np.sqrt(252)
            
            historical_data = {str(date): value for date, value in closing_prices.items()}
            
            return {
                'current_price': recent_close,
                'medium_term_average': medium_term_avg,
                'long_term_average': long_term_avg,
                'price_volatility': price_variability,
                'historical_prices': historical_data
            }
        except Exception as error:
            logging.error("Failed to obtain price data: %s", error)
            return {}

    def compute_financial_health(self, financial_data: dict) -> int:
        score = 0
 
        score += 1 if financial_data.get('asset_returns', 0) > 0 else 0
        score += 1 if financial_data.get('equity_returns', 0) > 0 else 0
        score += 1 if financial_data.get('net_margin', 0) > 0 else 0
        score += 1 if financial_data.get('liquidity_ratio', 0) > self.assessment_parameters['liquidity_ratio']['caution'] else 0
        score += 1 if financial_data.get('leverage_ratio', 0) < self.assessment_parameters['leverage_ratio']['caution'] else 0
        score += 1 if financial_data.get('asset_returns', 0) > financial_data.get('equity_returns', 0) * 0.5 else 0
        
        return score

    def estimate_fraud_probability(self, financial_data: dict) -> float:
        fraud_score = (
            -4.84 + 0.92 * self.fraud_detection_weights['receivables_index'] 
            + 0.528 * self.fraud_detection_weights['margin_index'] 
            + 0.404 * self.fraud_detection_weights['asset_quality'] 
            + 0.892 * self.fraud_detection_weights['growth_index'] 
            + 0.115 * self.fraud_detection_weights['depreciation_rate'] 
            - 0.172 * self.fraud_detection_weights['expense_index'] 
            + 4.679 * self.fraud_detection_weights['accruals_ratio'] 
            - 0.327 * self.fraud_detection_weights['leverage_change']
        )
        
        return fraud_score

    def analyze_media_sentiment(self, organization: str) -> dict:
        news_endpoint = f"https://newsapi.org/v2/everything?q={organization}&apiKey={self.news_api}&language=en&sortBy=publishedAt"
        try:
            news_response = requests.get(news_endpoint, timeout=10)
            articles_data = news_response.json()
            
            favorable = 0
            unfavorable = 0
            risk_terms = ['fraud', 'misconduct', 'probe', 'regulatory', 'irregularity']
            
            for story in articles_data.get('articles', [])[:50]:
                content = f"{story.get('title', '')} {story.get('description', '')}".lower()
                if any(term in content for term in ['positive', 'growth', 'bullish']):
                    favorable += 1
                if any(term in content for term in ['negative', 'decline', 'bearish']):
                    unfavorable += 1
            
            risk_mentions = sum(
                1 for story in articles_data.get('articles', []) 
                if any(term in f"{story.get('title', '')} {story.get('description', '')}".lower() 
                      for term in risk_terms)
            )
            
            return {
                'favorable_news': favorable,
                'unfavorable_news': unfavorable,
                'risk_references': risk_mentions,
                'articles_analyzed': len(articles_data.get('articles', []))
            }
        except Exception as error:
            logging.error("News analysis failed: %s", error)
            return {}

    def evaluate_company(self, ticker: str, company_name: str) -> dict:
        financial_metrics = self.retrieve_financial_statistics(ticker)
        market_data = self.obtain_price_history(ticker)
        media_analysis = self.analyze_media_sentiment(company_name)
        
        health_score = self.compute_financial_health(financial_metrics)
        fraud_score = self.estimate_fraud_probability(financial_metrics)
        
        current_value = market_data.get('current_price', 0)
        medium_avg = market_data.get('medium_term_average', 0)
        long_avg = market_data.get('long_term_average', 0)
        
        market_trend = "Neutral"
        if current_value > medium_avg > long_avg:
            market_trend = "Bullish"
        elif current_value < medium_avg < long_avg:
            market_trend = "Bearish"
        
        risk_category = "Low"
        if fraud_score > -1.78:
            risk_category = "Medium"
        if fraud_score > -1.0 or media_analysis.get('risk_references', 0) > 5:
            risk_category = "High"
        
        analysis_report = {
            'metadata': {
                'ticker_symbol': ticker,
                'organization': company_name,
                'evaluation_date': datetime.now().isoformat(),
                'data_providers': ['AlphaVantage', 'NewsAPI']
            },
            'financial_indicators': financial_metrics,
            'market_information': market_data,
            'media_evaluation': media_analysis,
            'risk_assessment': {
                'financial_stability': health_score,
                'fraud_probability': fraud_score,
                'risk_classification': risk_category,
                'market_trend': market_trend
            },
            'risk_notifications': self.generate_risk_alerts(financial_metrics, health_score, fraud_score, media_analysis)
        }
        
        self.store_analysis(analysis_report)
        
        return analysis_report

    def generate_risk_alerts(self, financials: dict, health_score: int, fraud_score: float, media: dict) -> list:
        notifications = []
        
        if health_score <= self.assessment_parameters['financial_health_score']['critical']:
            notifications.append("Critical financial weakness detected")
        elif health_score <= self.assessment_parameters['financial_health_score']['caution']:
            notifications.append("Financial health concerns identified")
            
        if financials.get('liquidity_ratio', 0) < self.assessment_parameters['liquidity_ratio']['critical']:
            notifications.append("Severe liquidity risk present")
            
        if financials.get('leverage_ratio', 0) > self.assessment_parameters['leverage_ratio']['critical']:
            notifications.append("Excessive debt burden identified")
            
        if fraud_score > -1.78:
            notifications.append("Potential financial irregularities suggested")
        if fraud_score > -1.0:
            notifications.append("High probability of financial misrepresentation")
            
        if media.get('risk_references', 0) > 3:
            notifications.append(f"Multiple risk references in media ({media.get('risk_references')} mentions)")
        if media.get('unfavorable_news', 0) > media.get('favorable_news', 0) * 2:
            notifications.append("Predominantly negative media coverage")
            
        return notifications

    def store_analysis(self, report: dict) -> str:
        def serialize_dates(data):
            if isinstance(data, (datetime, pd.Timestamp)):
                return data.isoformat()
            if isinstance(data, dict):
                return {key: serialize_dates(value) for key, value in data.items()}
            if isinstance(data, (list, tuple)):
                return [serialize_dates(item) for item in data]
            return data
        
        serialized_report = serialize_dates(report)
        output_file = self.reports_folder / "stock_analysis.json"
        
        with open(output_file, 'w', encoding='utf-8') as output:
            json.dump(serialized_report, output, indent=2)
            
        return str(output_file)


class CompanyDataVisualizer:
    def __init__(self, file_path):
        self.data = self.load_company_data(file_path)
        self.company_name = self.data['metadata']['organization']
        self.ticker_symbol = self.data['metadata']['ticker_symbol']
        
    def load_company_data(self, file_path):
        with open(file_path, 'r') as file:
            return json.load(file)
    
    def create_financial_health_chart(self):
        figure, axes = plt.subplots(2, 2, figsize=(12, 10))
        figure.suptitle(f'{self.company_name} ({self.ticker_symbol}) Financial Health', y=1.02)
        
        profit_metrics = {
            'Net Margin': self.data['financial_indicators']['net_margin'] * 100,
            'ROA': self.data['financial_indicators']['asset_returns'] * 100,
            'ROE': self.data['financial_indicators']['equity_returns'] * 100
        }
        axes[0,0].bar(profit_metrics.keys(), profit_metrics.values(), color=['#4CAF50', '#2196F3', '#009688'])
        axes[0,0].set_title('Profitability Metrics (%)')
        axes[0,0].grid(axis='y', linestyle='--', alpha=0.7)
        
        company_valuation = {
            'P/E Ratio': self.data['financial_indicators']['valuation_ratio'],
            'Market Cap ($B)': self.data['financial_indicators']['company_valuation'] / 1e9
        }
        axes[0,1].bar(company_valuation.keys(), company_valuation.values(), color=['#FF9800', '#E91E63'])
        axes[0,1].set_title('Valuation Metrics')
        axes[0,1].grid(axis='y', linestyle='--', alpha=0.7)
        
        financial_ratios = {
            'Current Ratio': self.data['financial_indicators']['liquidity_ratio'],
            'Debt/Equity': self.data['financial_indicators']['leverage_ratio']
        }
        axes[1,0].bar(financial_ratios.keys(), financial_ratios.values(), color=['#9C27B0', '#607D8B'])
        axes[1,0].set_title('Financial Structure')
        axes[1,0].grid(axis='y', linestyle='--', alpha=0.7)
        
        stock_prices = {
            'Current': self.data['market_information']['current_price'],
            '52W High': self.data['financial_indicators']['annual_peak'],
            '52W Low': self.data['financial_indicators']['annual_trough']
        }
        axes[1,1].bar(stock_prices.keys(), stock_prices.values(), color=['#795548', '#F44336', '#3F51B5'])
        axes[1,1].set_title('Price Levels ($)')
        axes[1,1].grid(axis='y', linestyle='--', alpha=0.7)
        
        plt.tight_layout()
        plt.savefig('/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Stock_analyses_output/financial_health_dashboard.png', bbox_inches='tight')
        plt.close()
        
    def create_price_history_chart(self):
        plt.figure(figsize=(14, 7))
        
        price_history = pd.DataFrame.from_dict(
            self.data['market_information']['historical_prices'], 
            orient='index', 
            columns=['Price']
        )
        price_history.index = pd.to_datetime(price_history.index)
        price_history = price_history.sort_index()
        
        plt.plot(price_history.index, price_history['Price'], color='#1f77b4', linewidth=2, label='Daily Close')
        
        moving_average_50 = price_history.rolling(50).mean()
        moving_average_200 = price_history.rolling(200).mean()
        plt.plot(moving_average_50.index, moving_average_50['Price'], '--', color='#ff7f0e', linewidth=1.5, label='50-Day MA')
        plt.plot(moving_average_200.index, moving_average_200['Price'], '--', color='#2ca02c', linewidth=1.5, label='200-Day MA')
        
        current_stock_price = self.data['market_information']['current_price']
        plt.axhline(y=current_stock_price, color='r', linestyle='-', linewidth=1, label=f'Current: ${current_stock_price:.2f}')
        
        plt.title(f'{self.company_name} ({self.ticker_symbol}) Price History with Moving Averages')
        plt.xlabel('Date')
        plt.ylabel('Price ($)')
        plt.legend()
        plt.grid(True, linestyle='--', alpha=0.7)
        
        plt.tight_layout()
        plt.savefig('/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Stock_analyses_output/price_analysis.png')
        plt.close()
        
    def create_risk_assessment_chart(self):
        figure, axes = plt.subplots(1, 2, figsize=(14, 6))
        figure.suptitle(f'{self.company_name} Risk Analysis', y=1.05)
        
        risk_metrics = {
            'Financial Stability': self.data['risk_assessment']['financial_stability'],
            'Fraud Probability': self.data['risk_assessment']['fraud_probability']
        }
        risk_colors = ['#4CAF50' if score > 5 else '#FF9800' if score > 3 else '#F44336' for score in [risk_metrics['Financial Stability'], 0]]
        risk_colors[1] = '#F44336' if risk_metrics['Fraud Probability'] > -1.78 else '#FF9800' if risk_metrics['Fraud Probability'] > -2.5 else '#4CAF50'
        
        axes[0].barh(list(risk_metrics.keys()), list(risk_metrics.values()), color=risk_colors)
        axes[0].set_title('Risk Scores')
        axes[0].axvline(x=0, color='black', linestyle='-', linewidth=0.5)
        axes[0].grid(True, axis='x', linestyle='--', alpha=0.7)
        
        news_sentiment = {
            'Favorable': self.data['media_evaluation']['favorable_news'],
            'Neutral': self.data['media_evaluation']['articles_analyzed'] - self.data['media_evaluation']['favorable_news'] - self.data['media_evaluation']['unfavorable_news'],
            'Unfavorable': self.data['media_evaluation']['unfavorable_news'],
            'Risk Mentions': self.data['media_evaluation']['risk_references']
        }
        axes[1].bar(news_sentiment.keys(), news_sentiment.values(), 
                  color=['#4CAF50', '#FFC107', '#F44336', '#9E9E9E'])
        axes[1].set_title('Media Sentiment Analysis')
        axes[1].set_ylabel('Number of Articles')
        axes[1].grid(True, axis='y', linestyle='--', alpha=0.7)
        
        plt.tight_layout()
        plt.savefig('/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Stock_analyses_output/risk_analysis.png')
        plt.close()
        
    def create_summary_report(self):
        report = {
            "Company Overview": {
                "Name": self.company_name,
                "Ticker": self.ticker_symbol,
                "Analysis Date": self.data['metadata']['evaluation_date']
            },
            "Financial Summary": {
                "Market Cap ($B)": round(self.data['financial_indicators']['company_valuation'] / 1e9, 2),
                "P/E Ratio": round(self.data['financial_indicators']['valuation_ratio'], 2),
                "Net Margin (%)": round(self.data['financial_indicators']['net_margin'] * 100, 2),
                "Current Ratio": round(self.data['financial_indicators']['liquidity_ratio'], 2),
                "Debt/Equity": round(self.data['financial_indicators']['leverage_ratio'], 2)
            },
            "Market Performance": {
                "Current Price": round(self.data['market_information']['current_price'], 2),
                "50-Day MA": round(self.data['market_information']['medium_term_average'], 2),
                "200-Day MA": round(self.data['market_information']['long_term_average'], 2),
                "52-Week Range": f"{self.data['financial_indicators']['annual_trough']} - {self.data['financial_indicators']['annual_peak']}",
                "Volatility": round(self.data['market_information']['price_volatility'], 4)
            },
            "Risk Assessment": {
                "Financial Stability Score": self.data['risk_assessment']['financial_stability'],
                "Fraud Probability Score": round(self.data['risk_assessment']['fraud_probability'], 2),
                "Overall Risk Classification": self.data['risk_assessment']['risk_classification'],
                "Market Trend": self.data['risk_assessment']['market_trend'],
                "Risk Notifications": self.data['risk_notifications']
            },
            "Media Coverage": {
                "Total Articles Analyzed": self.data['media_evaluation']['articles_analyzed'],
                "Favorable Sentiment": self.data['media_evaluation']['favorable_news'],
                "Unfavorable Sentiment": self.data['media_evaluation']['unfavorable_news'],
                "Risk Mentions": self.data['media_evaluation']['risk_references']
            }
        }
        
        with open('/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Stock_analyses_output/company_summary_report.json', 'w') as output_file:
            json.dump(report, output_file, indent=2)
            
        return report
    
    def generate_all_reports(self):
        self.create_financial_health_chart()
        self.create_price_history_chart()
        self.create_risk_assessment_chart()
        final_report = self.create_summary_report()
        
        print("Generated the following files:")
        print("- financial_health_dashboard.png")
        print("- price_analysis.png")
        print("- risk_analysis.png")
        print("- company_summary_report.json")
        
        return final_report


def Financial_analysis_Agent(ticker_symbol,ticker_name):
    analyzer = FinancialRiskEvaluator(finance_api_key="SERP API KEY",news_api_key="NEWS API KEY")
    analysis_result = analyzer.evaluate_company(ticker_symbol, ticker_name)
    visualizer = CompanyDataVisualizer('/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Stock_analyses_output/stock_analysis.json')
    visualizer.generate_all_reports()


class CorporateLegalMonitor:
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.reports_directory = Path("/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Legal_analyses_output")
        self.reports_directory.mkdir(exist_ok=True)

        # Configuration for regional legal monitoring
        self.jurisdiction_settings = {
            "default": {
                "data_sources": ["reuters.com", "forbes.com", "bloomberg.com"],
                "search_terms": ["legal action", "lawsuit", "litigation"]
            },
            "US": {
                "data_sources": ["justice.gov", "sec.gov", "courtlistener.com"],
                "search_terms": ["class action", "regulatory action", "DOJ case"]
            },
            "EU": {
                "data_sources": ["curia.europa.eu", "europa.eu", "ft.com"],
                "search_terms": ["ECJ ruling", "EU regulation", "GDPR case"]
            }
        }

        self.multilingual_legal_terms = {
            "en": ["lawsuit", "litigation", "legal dispute"],
            "es": ["demanda", "pleito", "disputa legal"],
            "fr": ["procès", "litige", "conflit juridique"]
        }

    def _prepare_search_query(self, corporation: str, jurisdiction: str, period: str) -> str:
        config = self.jurisdiction_settings.get(jurisdiction, self.jurisdiction_settings["default"])
        sources = " OR site:".join(config["data_sources"])
        terms = " OR ".join(config["search_terms"])
        

        language = self._identify_primary_language(jurisdiction)
        if language in self.multilingual_legal_terms:
            terms += " OR " + " OR ".join(self.multilingual_legal_terms[language])
        
        time_constraint = self._calculate_time_filter(period)
        
        return f"{corporation} ({terms}) site:{sources} {time_constraint}"

    def _identify_primary_language(self, jurisdiction: str) -> str:
        language_mapping = {
            "US": "en", "UK": "en", "CA": "en",
            "DE": "de", "FR": "fr", "ES": "es"
        }
        return language_mapping.get(jurisdiction, "en")

    def _calculate_time_filter(self, period: str) -> str:
        period_mapping = {
            "last month": timedelta(days=30),
            "last year": timedelta(days=365),
            "last 5 years": timedelta(days=5*365)
        }
        if period in period_mapping:
            cutoff_date = datetime.now() - period_mapping[period]
            return f"after:{cutoff_date.strftime('%Y-%m-%d')}"
        return ""

    def _identify_legal_authorities(self, content: str) -> List[str]:
        identification_patterns = [
            r"(Supreme Court|High Court|District Court)",
            r"(ECJ|European Court of Justice)",
            r"(SEC|FCA|DOJ|Financial Conduct Authority)"
        ]
        
        found_entities = set()
        for pattern in identification_patterns:
            found_entities.update(re.findall(pattern, content))
        
        return list(found_entities)

    def _categorize_legal_matter(self, content: str) -> str:
        classification_criteria = {
            "Corporate Governance": ["board", "shareholder", "fiduciary"],
            "Regulatory Compliance": ["regulation", "compliance", "violation"],
            "Commercial Dispute": ["contract", "agreement", "breach"],
            "Employment Matter": ["employment", "discrimination", "wage"]
        }
        
        normalized_content = content.lower()
        for category, indicators in classification_criteria.items():
            if any(indicator in normalized_content for indicator in indicators):
                return category
        
        return "Uncategorized"

    def _evaluate_case_severity(self, content: str) -> str:
        normalized_content = content.lower()
        
        if any(term in normalized_content for term in ["criminal", "indictment"]):
            return "Critical"
        elif any(term in normalized_content for term in ["fraud", "penalty"]):
            return "High"
        elif any(term in normalized_content for term in ["dispute", "violation"]):
            return "Medium"
        return "Low"

    def _store_analysis_results(self, analysis_data: Dict) -> str:
        output_file = self.reports_directory / "legal_analysis.json"
        
        with open(output_file, 'w', encoding='utf-8') as output:
            json.dump(analysis_data, output, ensure_ascii=False, indent=2)
        
        return str(output_file)

    def investigate_legal_issues(self, corporation: str, jurisdiction: str = "US",period: str = "last year") -> Dict:

        search_query = self._prepare_search_query(corporation, jurisdiction, period)
        api_endpoint = f"https://serpapi.com/search.json?q={search_query}&api_key={self.api_key}"
        
        try:
            logging.info(f"Analyzing legal issues for {corporation} in {jurisdiction}")
            api_response = requests.get(api_endpoint)
            api_response.raise_for_status()
            response_data = api_response.json()
            
            identified_cases = []
            for result in response_data.get("organic_results", []):
                case_description = result.get("snippet", "")
                case_record = {
                    "case_title": result.get("title"),
                    "source_url": result.get("link"),
                    "origin": result.get("source"),
                    "date_recorded": result.get("date"),
                    "description": case_description,
                    "jurisdiction": jurisdiction,
                    "legal_bodies": self._identify_legal_authorities(case_description),
                    "case_category": self._categorize_legal_matter(case_description),
                    "risk_level": self._evaluate_case_severity(case_description)
                }
                identified_cases.append(case_record)
            
            analysis_results = {
                "case_analysis": {
                    "subject": corporation,
                    "region": jurisdiction,
                    "time_period": period,
                    "analysis_date": datetime.now().isoformat(),
                    "total_cases": len(identified_cases)
                },
                "case_details": identified_cases,
                "risk_breakdown": self._compile_risk_analysis(identified_cases)
            }

            output_path = self._store_analysis_results(analysis_results)
            
            return {
                "analysis_status": "completed",
                "results_path": output_path
            }
            
        except Exception as error:
            logging.error(f"Legal analysis failed: {error}")
            return {
                "analysis_status": "failed",
                "error_details": str(error)
            }

    def _compile_risk_analysis(self, cases: List[Dict]) -> Dict:
        risk_assessment = {
            "severity_distribution": {"Critical": 0, "High": 0, "Medium": 0, "Low": 0},
            "category_analysis": {},
            "authority_references": {}
        }
        
        for legal_case in cases:
            risk_assessment["severity_distribution"][legal_case["risk_level"]] += 1
            case_type = legal_case["case_category"]
            risk_assessment["category_analysis"][case_type] = risk_assessment["category_analysis"].get(case_type, 0) + 1
            for authority in legal_case["legal_bodies"]:
                risk_assessment["authority_references"][authority] = risk_assessment["authority_references"].get(authority, 0) + 1
        
        return risk_assessment
class LegalCaseAnalyzer:
    def __init__(self, file_path):
        self.case_data = self.load_case_file(file_path)
        self.company_name = self.case_data['case_analysis']['subject']
        self.legal_region = self.case_data['case_analysis']['region']
        
    def load_case_file(self, file_path):
        with open(file_path, 'r') as case_file:
            return json.load(case_file)
    
    def create_severity_chart(self):
        severity_levels = self.case_data['risk_breakdown']['severity_distribution']
        
        plt.figure(figsize=(10, 6))
        severity_colors = ['#FF5252', '#FF9800', '#FFEB3B', '#4CAF50']
        case_bars = plt.bar(severity_levels.keys(), severity_levels.values(), color=severity_colors)
        
        plt.title(f'{self.company_name} Legal Case Severity Distribution')
        plt.ylabel('Number of Cases')
        
        for bar in case_bars:
            bar_height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., bar_height,
                    f'{bar_height}',
                    ha='center', va='bottom')
        
        plt.tight_layout()
        plt.savefig('/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Legal_analyses_output/case_severity_distribution.png')
        plt.close()
        
    def create_category_chart(self):
        case_types = self.case_data['risk_breakdown']['category_analysis']
        
        plt.figure(figsize=(10, 6))
        highlight_uncategorized = [0.1 if cat == "Uncategorized" else 0 for cat in case_types.keys()]
        plt.pie(case_types.values(), labels=case_types.keys(), 
                autopct='%1.1f%%', startangle=90,
                colors=['#2196F3', '#FFC107', '#9C27B0', '#4CAF50'],
                explode=highlight_uncategorized)
        
        plt.title(f'{self.company_name} Legal Case Categories')
        plt.tight_layout()
        plt.savefig('/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Legal_analyses_output/case_categories.png')
        plt.close()
        
    def create_case_timeline(self):
        all_cases = self.case_data['case_details']
        formatted_dates = []
        
        for case in all_cases:
            try:
                parsed_date = datetime.strptime(case['date_recorded'], '%b %d, %Y')
                formatted_dates.append(parsed_date)
            except:
                continue
        
        if formatted_dates:
            plt.figure(figsize=(12, 4))
            plt.hist(formatted_dates, bins=12, color='#3F51B5', edgecolor='black')
            plt.title(f'{self.company_name} Legal Case Timeline')
            plt.xlabel('Date')
            plt.ylabel('Number of Cases')
            plt.grid(axis='y', linestyle='--', alpha=0.7)
            plt.tight_layout()
            plt.savefig('/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Legal_analyses_output/case_timeline.png')
            plt.close()
    
    def create_risk_matrix(self):
        case_records = self.case_data['case_details']
        risk_details = []
        
        for case in case_records:
            risk_details.append({
                'Title': case['case_title'],
                'Date': case['date_recorded'],
                'Risk': case['risk_level'],
                'Category': case['case_category'],
                'Source': case['origin']
            })
        
        risk_df = pd.DataFrame(risk_details)
        risk_df['Risk_Score'] = risk_df['Risk'].map({'Critical': 4, 'High': 3, 'Medium': 2, 'Low': 1})
        
        plt.figure(figsize=(12, 6))
        risk_plot = plt.scatter(
            x=pd.to_datetime(risk_df['Date']),
            y=risk_df['Risk_Score'],
            c=risk_df['Risk_Score'],
            cmap='RdYlGn_r',
            s=100,
            alpha=0.7
        )
        
        plt.title(f'{self.company_name} Legal Risk Matrix')
        plt.xlabel('Date')
        plt.ylabel('Risk Level')
        plt.yticks([1, 2, 3, 4], ['Low', 'Medium', 'High', 'Critical'])
        plt.colorbar(risk_plot, label='Risk Severity')
        plt.grid(True, linestyle='--', alpha=0.5)
        plt.tight_layout()
        plt.savefig('/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Legal_analyses_output/risk_matrix.png')
        plt.close()
        
    def create_case_summary(self):
        report_summary = {
            "Company": self.company_name,
            "Region": self.legal_region,
            "AnalysisPeriod": self.case_data['case_analysis']['time_period'],
            "TotalCases": self.case_data['case_analysis']['total_cases'],
            "RiskBreakdown": self.case_data['risk_breakdown']['severity_distribution'],
            "CaseCategories": self.case_data['risk_breakdown']['category_analysis'],
            "NotableCases": [],
            "KeyFindings": {
                "HighestRiskCategory": max(
                    self.case_data['risk_breakdown']['category_analysis'].items(),
                    key=lambda x: x[1]
                )[0],
                "RiskConcentration": max(
                    self.case_data['risk_breakdown']['severity_distribution'].items(),
                    key=lambda x: x[1]
                )[0],
                "RecentHighRiskCases": []
            }
        }
        
        for case in self.case_data['case_details']:
            if case['risk_level'] in ['High', 'Critical']:
                report_summary['NotableCases'].append({
                    'Title': case['case_title'],
                    'Date': case['date_recorded'],
                    'Risk': case['risk_level'],
                    'Source': case['origin']
                })
                
                try:
                    case_date = datetime.strptime(case['date_recorded'], '%b %d, %Y')
                    if (datetime.now() - case_date).days < 180:
                        report_summary['KeyFindings']['RecentHighRiskCases'].append({
                            'Title': case['case_title'],
                            'Date': case['date_recorded']
                        })
                except:
                    continue
            
        return report_summary
    
    def generate_all_legal_reports(self):
        self.create_severity_chart()
        self.create_category_chart()
        self.create_case_timeline()
        self.create_risk_matrix()
        final_report = self.create_case_summary()
        
        print("Generated legal analysis files:")
        print("- case_severity_distribution.png")
        print("- case_categories.png")
        print("- case_timeline.png")
        print("- risk_matrix.png")
        
        return final_report




def Legal_analyisi_Agent(corporation,jurisdiction,period):
    legal_analyzer = CorporateLegalMonitor(api_key="GROK API KEY")
    analysis_output = legal_analyzer.investigate_legal_issues(corporation,jurisdiction,period)
    if analysis_output["analysis_status"] == "completed":
        print(f"Analysis saved to: {analysis_output['results_path']}")
    case_analyzer = LegalCaseAnalyzer('/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Legal_analyses_output/legal_analysis.json')
    case_analyzer.generate_all_legal_reports()


class CompanyBackgroundAnalyzer:
    def __init__(self, api_key: str):
        self.base_url = "https://api.opencorporates.com/v0.4"
        self.api_key = api_key
        self.reports_directory = Path("/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Company_Background_analysis_output")
        self.reports_directory.mkdir(exist_ok=True)
        self.data = {
            "metadata": {
                "created_at": datetime.now().isoformat(),
                "tool_version": "1.2",
                "data_sources": ["OpenCorporates"]
            }
        }
        self.logger = self._setup_logger()
        self.request_delay = 0.7 

    def _setup_logger(self):
        logger = logging.getLogger(__name__)
        logger.setLevel(logging.INFO)
        handler = logging.StreamHandler()
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        handler.setFormatter(formatter)
        logger.addHandler(handler)
        return logger

    def _make_api_call(self, endpoint: str, params: Optional[Dict] = None) -> Dict:
        if params is None:
            params = {}
        params['api_token'] = self.api_key
        
        try:
            response = requests.get(
                f"{self.base_url}/{endpoint}",
                params=params,
                timeout=15
            )
            response.raise_for_status()
            
            result = response.json()
            self._store_api_call(endpoint, params, response.status_code, result)
            return result
            
        except requests.exceptions.RequestException as e:
            self.logger.error(f"API request failed for {endpoint}: {str(e)}")
            self._store_error(endpoint, str(e))
            return {"error": str(e)}

    def _store_api_call(self, endpoint: str, params: Dict, status_code: int, result: Dict):
        if "api_calls" not in self.data:
            self.data["api_calls"] = []
            
        self.data["api_calls"].append({
            "endpoint": endpoint,
            "parameters": params,
            "status_code": status_code,
            "timestamp": datetime.now().isoformat(),
            "response": result
        })

    def _store_error(self, endpoint: str, error_msg: str):
        if "errors" not in self.data:
            self.data["errors"] = []
            
        self.data["errors"].append({
            "endpoint": endpoint,
            "error": error_msg,
            "timestamp": datetime.now().isoformat()
        })

    def search_companies(self, company_name: str, jurisdiction: Optional[str] = None, 
                        limit: int = 5) -> List[Dict]:
        params = {
            'q': company_name,
            'per_page': limit,
            'sparse': 'true'
        }
        if jurisdiction:
            params['jurisdiction_code'] = jurisdiction.lower()
            
        result = self._make_api_call("companies/search", params)
        return result.get('results', {}).get('companies', [])

    def get_company_full_profile(self, jurisdiction_code: str, company_number: str) -> Dict:
        endpoints = [
            ('company_details', f"companies/{jurisdiction_code.lower()}/{company_number}"),
            ('officers', f"companies/{jurisdiction_code.lower()}/{company_number}/officers"),
            ('filings', f"companies/{jurisdiction_code.lower()}/{company_number}/filings"),
            ('ownership', f"companies/{jurisdiction_code.lower()}/{company_number}/owners"),
            ('subsidiaries', f"companies/{jurisdiction_code.lower()}/{company_number}/subsidiaries")
        ]
        
        profile = {}
        for name, endpoint in endpoints:
            profile[name] = self._make_api_call(endpoint)
            time.sleep(self.request_delay)
            
        return profile

    def analyze_company_network(self, jurisdiction_code: str, company_number: str, 
                              depth: int = 1) -> Dict:
        network = {}
        current_depth = 0
        
        while current_depth < depth:
            endpoint = f"companies/{jurisdiction_code.lower()}/{company_number}/network"
            result = self._make_api_call(endpoint)
            
            if "error" in result:
                break
                
            network[f"depth_{current_depth}"] = result.get('results', {})
            current_depth += 1
            time.sleep(self.request_delay)
            
        return network

    def save_analysis(self, filename: str = "/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Company_Background_analysis_output/company_background_analysis.json") -> Tuple[bool, str]:
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(self.data, f, indent=2, ensure_ascii=False)
            self.logger.info(f"Analysis successfully saved to {filename}")
            return True, filename
        except Exception as e:
            self.logger.error(f"Failed to save analysis: {str(e)}")
            return False, str(e)

    def enhance_with_external_data(self, company_data: Dict) -> Dict:
        company_data["external_data"] = {
            "last_updated": datetime.now().isoformat(),
            "notes": "External data integration not configured"
        }
        return company_data


def company_Background_anslysis_Agent(company_name,region):
    analyzer = CompanyBackgroundAnalyzer("API KEY")
    companies = analyzer.search_companies(company_name, region, limit=3)
    
    if not companies:
        print("No companies found matching search criteria")
        return
    company = companies[0]['company']
    print(f"Analyzing: {company['name']} ({company['jurisdiction_code']})")
    
    full_profile = analyzer.get_company_full_profile(
        company['jurisdiction_code'],
        company['company_number']
    )
    
    network_analysis = analyzer.analyze_company_network(
        company['jurisdiction_code'],
        company['company_number'],
        depth=2
    )
    
    analyzer.data["company_profile"] = full_profile
    analyzer.data["network_analysis"] = network_analysis
    
    success, filename = analyzer.save_analysis()
    if success:
        print(f"Company background analysis saved to {filename}")
    else:
        print(f"Failed to save analysis: {filename}")



from langchain_ollama import OllamaLLM
import json
from typing import Dict, List
import textwrap
import os




def chunk_data(data: Dict, max_chars: int = 1500) -> List[str]:
    """Smart JSON chunking that preserves structure"""
    if isinstance(data, dict):
        chunks = []
        current_chunk = {}
        current_size = 0
        
        for key, value in data.items():
            item_str = json.dumps({key: value})
            if current_size + len(item_str) > max_chars and current_chunk:
                chunks.append(json.dumps(current_chunk))
                current_chunk = {}
                current_size = 0
                
            current_chunk[key] = value
            current_size += len(item_str)
            
        if current_chunk:
            chunks.append(json.dumps(current_chunk))
        return chunks
    return textwrap.wrap(json.dumps(data), width=max_chars)

def analyze_with_llama(data: Dict, analysis_type: str) -> str:
    llama = OllamaLLM(
        model="llama3.3",
        temperature=0.2,
        top_p=0.9,
        repeat_penalty=1.1,
        num_ctx=4096
    )
    """Enhanced analysis that combines chunks into unified analysis"""
    chunks = chunk_data(data)
    context = ""
    full_analysis = ""
    
    system_prompts = {
        "financial": """You are a financial risk analyst. Analyze and combine all chunks to create:
1. Unified financial assessment
2. Consolidated risk score (0-100)
3. Integrated key metrics analysis
4. Complete fraud risk evaluation""",
        
        "corporate": """You are a corporate structure analyst. Synthesize all chunks into:
1. Comprehensive entity status report
2. Complete ownership structure
3. Final data quality assessment
4. Unified complexity score""",
        
        "legal": """You are a legal risk specialist. Combine all chunks to produce:
1. Complete case volume analysis
2. Final severity assessment
3. Integrated jurisdictional risk
4. Consolidated legal exposure score"""
    }

    all_chunks = []
    for i, chunk in enumerate(chunks, 1):
        prompt = f"""SYSTEM: {system_prompts[analysis_type]}
        
CONTEXT FROM PREVIOUS ANALYSIS:
{context[-800:] if context else 'No prior context'}

NEW DATA CHUNK ({i}/{len(chunks)}):
{chunk}

INSTRUCTIONS:
- Analyze this chunk in context of complete analysis
- Prepare to combine with other chunks
- Identify connections to previous data"""
        
        response = llama.invoke(prompt)
        all_chunks.append(response)
        context = f"{context}\n{response}"[-1000:]
    
    chunk_separator = "\n\n".join(all_chunks)
    combine_prompt = f"""SYSTEM: Combine all analysis chunks into unified {analysis_type} report:

ANALYSIS CHUNKS:
{chunk_separator}

INSTRUCTIONS:
1. Remove duplicate information
2. Organize by key risk categories
3. Provide final {analysis_type} risk score (0-100)
4. Highlight most critical 3 findings
5. Format in markdown with headings"""
    
    full_analysis = llama.invoke(combine_prompt)
    return full_analysis

def generate_executive_summary(analyses: Dict[str, str]) -> str:
    llama = OllamaLLM(
        model="llama3.3",
        temperature=0.2,
        top_p=0.9,
        repeat_penalty=1.1,
        num_ctx=4096
    )
    """Create polished executive summary from unified analyses"""
    summary_prompt = f"""Create executive summary from these complete analyses:
    
FINANCIAL ANALYSIS:
{analyses.get('financial', 'No data')}

CORPORATE ANALYSIS:
{analyses.get('corporate', 'No data')}

LEGAL ANALYSIS:  
{analyses.get('legal', 'No data')}

Structure as:
1. Overall Risk Rating (Low/Medium/High)
2. Key Findings (3-5 bullet points)
3. Critical Risks (Top 3)
4. Recommended Actions
5. Final Risk Score (0-100)"""
    
    return llama.invoke(summary_prompt)

def analyze_all_files(file_paths: Dict[str, str]) -> Dict[str, str]:
    """Process all JSON files with unified analysis"""

    analyses = {}
    print("Starting unified analysis...\n")
    
    for file_type, path in file_paths.items():
        if not os.path.exists(path):
            print(f" File not found: {path}")
            continue
            
        print(f"Analyzing and combining {file_type} data...")
        with open(path, 'r') as f:
            data = json.load(f)
        
        analyses[file_type] = analyze_with_llama(data, file_type)
        print(f" Unified {file_type.upper()} analysis completed\n")
    
    return analyses

# def save_report(content: str, filename: str = None):
#     """Save report with timestamp"""
#     if not filename:
#         timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#         filename = f"Infosys_Unified_Risk_Report_{timestamp}.md"
    
#     with open(filename, 'w') as f:
#         f.write(content)
#     return filename

def save_report(content: str, filename: str = None, file_path: str = "/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/"):
    """Save report with timestamp to a specific directory"""
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Infosys_Unified_Risk_Report.md"
    
    full_path = os.path.join(file_path, filename)
    
    os.makedirs(file_path, exist_ok=True)  # Ensure the directory exists
    with open(full_path, 'w') as f:
        f.write(content)
    
    return full_path


def get_all_image_paths(folder_path: str, image_extensions=None):
    """Returns a list of all image file paths in the folder (including subfolders)."""
    if image_extensions is None:
        image_extensions = ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.webp']

    image_paths = []

    for root, dirs, files in os.walk(folder_path):
        for file in files:
            ext = os.path.splitext(file)[1].lower()
            if ext in image_extensions:
                image_paths.append(os.path.join(root, file))

    return image_paths


def image_markdown_if_exists(image_path: str, title: str = "Image"):
    base_url = "https://8000-01jrcj02963afp5r8dxtpsxdqz.cloudspaces.litng.ai"

    if os.path.exists(image_path):
        full_url = f"{base_url}{image_path}"
        return f"## {title}\n\n![{title}]({full_url})\n"
    return ""


def create_legal_cases_table(json_data):
    headers = ["Case Title", "Date", "Jurisdiction", "Risk Level", "Source"]
    table_md = "| " + " | ".join(headers) + " |\n"
    table_md += "| " + " | ".join(["---"] * len(headers)) + " |\n"

    for entry in json_data.get("case_details", []):
        source_text = entry.get("origin", "")
        source_url = entry.get("source_url", "")
        source_link = f'[{source_text}]({source_url})' if source_text and source_url else source_text

        row = [
            entry.get("case_title", ""),
            entry.get("date_recorded", "") or "N/A",
            entry.get("jurisdiction", ""),
            entry.get("risk_level", ""),
            source_link
        ]
        table_md += "| " + " | ".join(row) + " |\n"

    return table_md

async def convert_html_to_pdf_async(html_path, pdf_path):
    process = await asyncio.create_subprocess_exec(
        'node',
        '/teamspace/studios/this_studio/Uvarajan/react_templates/react-app/src/convertToPd.js',
        html_path,
        pdf_path,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE
    )
    stdout, stderr = await process.communicate()
    
    if process.returncode != 0:
        print("Error:", stderr.decode())
        return False
    return True


def insert_watermark_div_in_html(file_path, output_file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        html_content = file.read()
    watermark_div = """
    <div class="watermark1" style="
        position: fixed;
        top: 0;
        left: 0;
        width: 100%;
        height: 100vh;
        background-image: url('data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAVIAAAEOCAYAAADbmZ9jAAAAAXNSR0IArs4c6QAAIABJREFUeAHtneuTHcd53vkn8E/gl3x0RXEqTqJyHMapWEmlpFIcxyXFrpiW48RxuRToS2xHcokVSY6UDyalWKrIcUwmqiQlWTIhOzZJSSRA4rbYBbC4L4DdxWJxW9yBXYA3gVRHv97zHPYOzjlz5vTs7pwzz1RNza27z/Qz07953+535jz22IhTCOHJEMKOEMILIYSlEMLd4MkKWAErMD4KwKzZDsOehmkj4rBathDC4z/W6LOG5vjcKT5TK2AFKimAUfhcCOGJanQcInUHoM9UOh0ntgJWwAqMtwL1AbXjvtttH+8bwmdvBazAaApgoT41hL3ZP0kIwVboaOI7lxWwApOlwDP9SdnnSMeV3zVZOrg2VsAKWIEsBRiYerwPNh/d3RnJyvpFZ7YCVsAKTKACux4lZo89ducn8NK7SlbACtSpwGA3n07VOn/NZVkBK2AFJlSBHT3s0MceI26qE1g/ofV2tayAFbACtSlAJNOj/aU/Bunztf2EC7ICVsAKTL4CG138jjU6+dV2Da2AFbAC9SrwvlVqa7ReZV2aFbACrVHg6W5fqftGW3PRXVErYAXqVeBuBGnnK071Fu3SrIAVsALtUeBJRuv5fJQnK2AFrIAVGE2BHYB052h5ncsKWAErYAVgKCDl/VFPVsAKWAErMJoCS4DUn8gbTTznsgJWwAqgwF1A6skKWAErYAUyFDBIM8RzVitgBawAChikvg+sgBWwApkKGKSZAjq7FbACVsAg9T1gBayAFchUwCDNFNDZrYAVsAIGqe8BK2AFrECmAgZppoDObgWsgBUwSH0PWAErYAUyFTBIMwV0ditgBayAQep7wApYASuQqYBBmimgs1sBK2AFDFLfA1bACliBTAUM0kwBnd0KWAErYJD6HrACVsAKZCpgkGYK6OxWwApYAYPU94AVsAJWIFMBgzRTQGe3AlbAChikvgesgBWwApkKGKSZAjq7FbACVsAg9T1gBayAFchUwCDNFNDZrYAVsAIGqe8BK2AFrECmAgZppoDObgWsgBUwSH0PWAErYAUyFTBIMwV0ditgBayAQep7wApYASuQqYBBmimgs1sBK2AFDFLfA1bACliBTAUM0kwBnd0KWAErYJD6HrACVsAKZCpgkGYK6OxWwApYAYPU94AVsAJWIFMBgzRTQGe3AlbAChikvgesgBWwApkKGKSZAjq7FbACVsAg9T1gBayAFchUwCDNFNDZrYAVsAIGqe8BK2AFrECmAgZppoDObgWsgBUwSH0PWAErYAUyFTBIMwV0ditgBayAQep7wApYASuQqYBBmimgs1sBK2AFDFLfA1bACliBTAUM0kwBnd0KWAErYJD6HrACVsAKZCpgkGYK6OxWwApYAYPU94AVsAJWIFMBgzRTQGe3AlbAChikvgesgBWwApkKGKSZAjq7FbACVsAg9T1gBayAFchUwCDNFNDZrYAVsAIGqe8BK2AFrECmAgZppoDObgWsgBUwSH0PWAErYAUyFTBIMwV0ditgBayAQep7wApYASuQqYBBmimgs1sBK2AFDFLfA1bACliBTAUM0kwBnd0KWAErYJD6HrACVsAKZCpgkGYK6OxWwApYAYPU94AVsAJWIFOB1oB09e03w8ra3XD25tUwe2UpHLq0GKaW51s5U/fjKxejFpfu3Q5o48kKWIHRFZhokD58771w+d7tVkNz2IfF7JULYfHW9fD2w4ej303OaQVaqsBEghSAXrhzc4O1OXPpfDhxYyXM3b4VFu+vhQtvPAjLb77Zypm6L6zdC2fu3o6aoE0K3AUDtaU4cLVHVWDiQLqyeneDBXrs2pVw7t7dVgKzyoMCsKKVgIqFeuPBWvjRj3406r3lfFagNQpMFEiXEiv0yNXlcP7+mgFa0epGsxSoy3duGqatwYErOqoCYw9SLCZc+dPXLkdr6uDFhXD69g0DtCJAi9brqdvXA1pioTIw9fC9d0e9x5zPCky8AmMNUiDKTEOnwdPXh4tahIK3R+sLxjpV/ykPKuk98a3CFbQCFRUYW5CqUS/dvtGFqF350YA56EGTwpQBPHT3ZAWswEYFxhKkguj1+/e67rwhWj9EBVisfLn5V1fv2DLd2Ia8ZQXC2IEUiL733nvhzR++HQPrcendJ7p5EBVM6TNd7z5ZDO+8+9CWqeFhBRIFxgqkskQB6cLNldiwGWFWY/dyc4F65Op6XzTdKVwDZrv6SWvyamsVGDuQ0njfeOetCFEsJLv0mwvP9OGE1mjO/PbDH4Z3333Xbn5r0eGKpwqMDUhljdJ4522NbpsVrhhT4ku5FrZK0+bk9bYqMBYgTSH68OHDcOTyUrSKFtdWtw0oqaXWpvX51btR+5lLi+GHP1y3SoGpJyvQZgXGBqQ0ViB6a219pJ74xjYBrEl1VWzprbXVLkzdV9pmjLjujQcpDVQQfeedd8LSrfXRYz5A0iS4tOlcjt+4Gq3SS7dvBq4JDzi7+IZJmxVoNEiLLv1bb70VTl+7FBvx2bu3DdLM10BHhf/cnVvxGpxeuRTefPPNCFIPPLUZI65740GKpUMjxfJ58OBBmHX/6LY/QDR6f/jS+QjSolVqN99gaZsCYwFSBjWwfO7fvx+mOx/SuPDGG9sOlFEtunHPt/TGg/UBp4sL8ZrIKpV7b5C2DSOub+NBKmv0jTfeCKurq7EBE8c47jAa9/NXPOm9e/cC10ZWqbpj3LSsQJsUaCxI1SAZyHj77bejW3/nzp0tAenp6zfDH33rz8NTv/Ufw0//3MfCX/uJnx2L+R///L+O5/zf//QvNv1BI5ByTfAUuEYedGoTOlzXVIFGgxRrVG792tpauHVrfZBjsyxSugv+74uvhr/xwQ+PBTgHAf7v/dzHw2YCVSC9fft24NowEKi4Urv2aRPzehsUaDxIcRlxHXEhb9xY/2TeZoH0t7/w5bEHaBGun/vKH2+KdSqQ3rx5M14brhFWqUbv29B4XEcrIAUaCVIsmjR2lNH6u3fvhuvX12NINwOk/+nLfzxxEBVUeUDU3ScrkHJNuDZcI7n36pbRTealFZh0BRoLUhpj2j+KC7mysv7Fp7pB+kff+u7EQlQw/d9/9UpIIx3S9VEgK5ByTdxPOumYcP3KFGg0SOlzo++NwQxAevXq+hs1dYEUmDCP04CSwFh1Sb/vyWs3Yn1zIQp4U5AW+0nxJjxZgTYp0FiQ0hgFUsKeGGi6cuX9vwsexYpK8wCTpQcPwte/OfnWqKD73765M9ZZD5BUj6rrAikPN0DKNfKAU5vQ4bqmCjQapAw0EexNI2VQ4/Ll9X8KzbVIBdHz9++HX/nN3514t14g/ZXf/J1AnXmA5MJUIOXhxkOOa8S14pp5wCltYl5vgwKNBCnCF0OfGLGvA6QpRBdXV8OHfv7XWgPSn/5HHwvUOYVpVUtU6VOQ8pADpArMN0jbgA7XMVVgLECq0KdLl9Y/WDKqRSorDJAAlPm7d1sDUVml1Pn82lq2VZqClIdc+oYTIPVkBdqkQCNByoi9LNI0hjTXIpU1uri2Fhbu3Qvn7txpHUjP3r4d6w5MeaCgiazMKkuBlGuSWqQOym8TPlxXKdBokKZ9pFg9dVikskaB6OkbN1oH0tPXrwdgikVOX6n6S6tAlLS9QEosqftI1bS8bJMCjQWpRu0ZwKjDtS9ao2du3gwnr11rHUhPrKwE6o5FjmUOSC++9VZlq7QIUrn2WKRcO7wKT1agLQo0EqSIj2uPdSPXHvcxxyIFpKk1OnfjRjh25UrrQHr86tVw6vr12D8sq3QU9z4FKd4Cg02ySHmRognT0tJSeP7550vnXbt2ZZ/u7OzswN/h7a86J8obpm45adAvd9J57tixI3z0ox8NH/jAB8ITTzwR5yeffDLu49jOnTvjG3K5v7dd+RsJUvWRpiDNde27FunqanRtT127Fo5evtw6kFJn6n7m1q31vtJOP2lVmAqkPNx6DTY1wSJ96qmnwmOPPVY6P/7449nt7+mnnx74O3XAOj1JIDdM3XLSAOFRp927dwdAWfX3gS15x21qDUhxYTXIhGuLi3tkeXnLQPrkhz4envvGt8N/+PQXw09+8CM9f5c039n5YvjOzpcC6xppr3N5eHk5nFxZCXM3b65bpWtrIw04CaQMNqUg1YdLmtAQsHyGbci5oDNI16/4qAAtXicgXIdFvFX3YaNBSn9bXX2kAinhP7j1uLiHlpY2BVZF8P3kBz8cLl1e6V7Tqemjj/wu4EzTsN4PuMXyq2wfvnAh1h0NGHCTez/qYJMs0jSOlD5S5u2caNDFxjloG+s1ZzJIQyjTYJD+/Y49++yzOZdly/I2FqQ0RLn2NNIc1z5164EH1hgu7sz5848ArQqUhk2LhVmcsE7T/FPTs8Uk0YJN09SxTp3pG6afVCAdJQxKFikgpf9ag00atd9ukA7r1qsB497n9GOWQSTX4i3eHE1z7avqLd2HWaJt06exACmNlM+15Qw2RYt0dTXCA4gA0umFhQ0wqwNUxTI+/6U/7HsP/NKvfir+/t/84EfC2tr9nul+45O/V+s5Ti8uxroT+tWNKR0hnlQgLcaRAlIgut19pFXcejXmHNi1GaQMFknDzVo23TJtJEghigLy+RBGbvgTFqlG7BlkAaSzly5tOkhx1wdNqfv+hS99tWfS1bX7tfaXAlLqjgaKJ0WbHNdeFimj9mlA/nbBtKpbr8ZPv9yoU1tBWlZvaVvHksiIpk6NBCkNEKtGfaS4XLj2Fy9e7AaCV234KUiJHwUmBzfRIi32i/a7AdL+0u+/srdnsjRN0eKtuk2dqTsaFC3SKiP3qUWaDjalIO1ZmS3YOaqbmePelwElx9rtJVmZa5/b59vrN4v7ys5B8ERXrFY0SLtPWGffsNcr50FXPPe6txsLUlmkiiPNAan6SAlCBx6E/xwByvPztbrNKdSe+1/fHvpafe6Lf9h18S9fudYzX7FPNf2tKusRpBcv9gRplYdTP5Cqj3S7rFHEo+GqEVddjhry00aQDhPeRNzoMKPvpBmmO6buB1LPxjbCzkaDVINNdbn2gDS69likmwhSQpyqTuov/aVPfKpvVqWpAs5i2ujad0CaxpJWgShpBVKN2nONFJDPQ3C7QEpg9yB40tc2CLSjWj1tA+kw3SfEhFaZsFAB76DrN+r1qXIeo6RtNEjl2ueO2tPwGWyKFil9pJsIUvpF6desOg3TX5qmKQJy2G2scB4iuPYpSKu49f1A2oTP6JW5iVg+ZWlS93PY69g2kJZpiHU5io7DdBeMUu6w13HUdI0GqSxSgTSnj1Sj9rj2QGSzXHtgN+r0vVf2drsapmaO9iyGUKphodkrnfpI4zv3hbebqlilskg1ao9FSswvAfnbNWpPAxvGmhnGau0p/oCdbQJpmc5cg1G7SJC4DNJcv6ZNjQdpHQH56iMl+BwrjDd7omt/7lwWlIqgqtIv2u9GkPv+5If+Zc+QKKzdnEB9LFIeIqlFulRD+BMg3W6L9LnnnhsIUjVuQFC3e98mkJY9iLBGc6ay8tG6aVMjQYpIxcEm4kixfmQJVbGeomt//3507SNI5drXCNJR+kV73QzpoNJXvtb7XecP/8Kvj/wAiK59Z9Q+de2x2KtoquugPtL0zabt6iOlT26QRZoOepQNlFR1H9sE0rK4UT2wet3fw+xDex6K/eYmDjg1FqQKf9Kofe7XnxT+1HXtl5fDwZpG7Yuvdw5zs/RLk7ru/UD6y51A/qJFPMy2Ru3TD5cA0VH7SOXaA1K8h+0atS/rWysOUpRZr1UDwNsE0rKHUJPjPfu1u9z9jQQpI76pRUojBaSj9pECiSJI6xy1z+kXLV7A3/7Ml6K1OSgONde1j4NNfJf01q344RK0yQFpGkcKSLejj7QMjEUrqW73vmkgBXbUedR5kEU+KEyJLpM2To0FqSxSrBwNNo3q2sc+0s7/NMkiBSZ1WKRf7uN+j3IzEUMqSPZ6P58yD/b44MkwlqjSYJGmfaT6aEkOSHnIbbdrX2YlpW69rk1ZnkEwURlaNg2kg7o4hjl24cIFVe2R5aD8hC+1cWosSHtZpKO+ay+LdKEz2KTwp1yQ/sYnP1PrPcMAE8Ab9H6+0giMVZddkNZokeoV0e0abKrq1uui4b4PggJwHHZqC0h5uAzSjIdTG6dGgzQNf8rpI+2CtBOQH8OfaugjnZqp793fr3z1f0aIDno///Nf+urIg0wCbheknThSLNIc156HW2qRbscrolXdejX0OqHQFpCO+tCS5pO6bDRIFZBPaE3OqL1ACjTk2tcRR9rvi01VbxbesQd0g/pF00EoQXGUpQab0vCnHJDS3VLsI93qUfsyF72XW69rVJZ3kIurMlgapOv/RGCLNL0rtnldg02ySAEpjXVU156wHg02KY4UkOa69pevjB58L4npF5W73q9fNO07HQWeaR7Cn44m4U+5FmkK0u0YtS+zkMr67Opy79sCUu5bu/Zqve8vG22RCqQabBoVpFik8RVRBeTX9GbT898Y/sMk70u+ce2XP7H+TdLN7BdNQYpFmoKU12brtEj1VyNb9a59mVtfFsZU5t4PG1xukK5bpMPqtbEVjP9Wo0GKa88ARh0gBRbFj5YAlRQyVdf7vX007G2xVf2iab169ZHmxpHKtcci3WqQln3kgvAfPrAxaB4UzoP1NUwAeNNAymuWmzWV6VUl2mGzznGry20sSBX+pID8XNc+WqTpYFNNX38a9LWmQRfz4Mz6/zZtRb9oL5CmAfn8KSBW+yhvNqWu/VaP2pe59YNc0CrHgGTZ1CaQlr1BNsyDp0zPcTveSJAiIgMWcu1z+0jl2muwiQ921BmQX9XF3+p+0SJI9WHn9BXRcQRpWf9mFVgOSjtMkHmbQFpW12EePINAyQNy0IsETQR140CqvrWiRZrzn01AInXt9fWn3MGmFFCn5xYG3Rsbjv27zv8wbVW/aHqecdSevxrpfCGff1XFWq9ijZI2fde+GEf68OHDLfkeaZmLOQiOVY+VNd4yuJTl33CDDLFRZo1vpmtPXQbpx4Mnx70v6/fOBfUQ8lZO0jiQUoNeo/a5IJVrr/AnrDJGsFPI5KwP21+6Hf2iab3q/voTA4B0u2z1m028zz2oMdd9jA91DJraBFJ0GPT1LLTPgV1ZSFrdD6VB13XYY40FKRYprj0DGHW49hvCnzr/2VQnSIEVluagCZeedFvdL5qCNFqkfEav82ZTHaP2RYt0K+JIy75AVDdIy9z7toG0rL7ohdVcdSqzRpsaFdBYkKqPVO/a12WR0i+oV0SnavyMnmDVr7807RcdlEbv2qu8upfpqP3czZsxkqHO8CceflsB0jK3ntF8LJsqc1mZgyyhMrAMylsVNqTfTtee3y8LG+NBhp5VYIqXUWbpbmaXxSjXQXkaC1JZpHWM2sfBph4fLanbIgV6/Ed9r/5SfdWJNHQDANbipMD8uuGZlqc+0rrfbMK136qA/LL/CxrVaimD4SD3vizvpIGUe3cYrwAwDhPLi35lEAXOVcBcbF+bud1IkFLhXiDN+YwefaSM2sc3mzbhw84prIr9pc994zuP9MUWw6bqeI8+PYd+65sFUrpftir8qeyvKAYBb1BjKrPyaOj9pjKQck6DRqIHHev11xpl54olPqjMYY6VfVcUq7TMilcXC+m4bkCV+jCzzr5hAEo5n/3sZ/vJv+37GwlSDTal79ozoDHqZ/T0iqgC8jVqvxkWqQAGTPnafWqJ6piWH/mFfxPT6O0m7d/MZTrYxMBbHX2k6WDTVny0pKzx5lh/o5ZdBlIBZZQl51ScykA6yu8U81Cnsgmti/k2Y7uXBmXntpXHGwvSYvhTTkC+4kiBBvA4df16/CZnneFPmwm/OsvuWqSdwaY637WXRcq1Uxhb3TfzZrn1Os8yIPbroyvLlwOXXhBpCkjRrWyAKKfu5KX+1LfJU2NBmg425Y7aY5Eq/Cl17dsK0vTDzrJIq8aSKo50q99sKnPr+4Fu2EZYZmHhhvaKkWwzSNH2mWee2RTLdBwgSv0bC9JefaQ5Hy15JPyppldE67QWt6KsdNS+DtdecaQ87LZisImGNcjCyXHrBdtRfqPtIEU7rMYy7QZdu+IxXkXt9dDSdWrSstEg1UdLci3S9M2m1LXfzD7SrYDiKL/R7SOtMY6UbhdueIF0s1z7sr/ppRHXMZWNRjOQU5wM0vcVYVAoB6jkreOB+P4Zbf5ao0Fa17v2qWsPSLuDTZsQRzoK3LYyDyCNf37X+UJ+Ha59MSB/s14R3Wy3Xs1tFPfeIJV660serPSd8tApWpq9tuky4fqOG0BV68aCdDP6SDeEP9X8iuhWwjDnt6YXF8MjHy2p+etPmwVS3bRejp8CAPKFF16Ir47y0GEm/AnYloVZjUNtGw3SYvhTTh/pI4NNly7FL+T/3Z/9F4/EeOaAqsl5/+GH/1VI+0i7X3/KAGnaR6pRex6CnqxAmxRoPEhpnLw1k/OKaDeOtBCQz6j9L/7qjtaA9Bc/sSN+qCVapCsrMRQM156HDP3I6DTsrFF7QJr++d1Wf9i5TY3VdW2uAo0EKXJh1aSDTTTW3DebgIbCn/Tnd//5a3/SGpB+8Wt/sg7Smj9awmDTVo3aN7cp+czarEAjQdrvzaYc1z5+jzT9X3tc+4WF8PL0TPjrf/ufTjxM/84/+Ofh9ZMnY50f6SO9f39kizSNI9WoPQ/BzQrIb3Njdd2bq0BjQVpnHGnXtdebTZ137acXFiJc/uB/fGPiQfr5L3897Dl1qgvS9K9Gcr7+JNcei1R9pAp/Mkyb2/B9ZvUq0FiQyrXHyqGR4trnWKT0A+qvRgh/wioTSH9w+HD49U/93sTC9JOf/v3wypEjG0CKBjEgf3U1q49UFulWf9i53mbg0qxAngKNBClVwqqpe9ReID19/fpGkB45El46eDD81u98buJg+mv//tPhxYMHwyuzs2Hv6dNB4U9okL7ZNOwgk9JpsEkgLVqkebelc1uB8VKgkSBVH6m+kI+1U8dHSwDpuTt3wukbN8Lxq1fDzPnzES6vzs6Gl6enw//bty985r98JfzU3/9nYw/Un/ipfxK+8F+/Huv08sxMePXo0bBvbi7W+diVK1EDQIomGrWvMnI/CKTuIx0vCPhs8xVoJEipFo2x+GbTqKP2WFHAgr5A/uyNL8MD0kNLS2H/mTNh97Fj4XszM+GvDhwIf75nT3hh167wu7//B+EjH/+34Wc+9LGxgerf+pmPxnPm3P/PX74Yvvv667FO3zt0KOw6diwcOHcuHL5wIdYdDXionF9bi9rI0hx2mYK0+GaT+kjzb0+XYAXGQ4FGglQWqcKfci1S4IC1xf+3C6T8JfOR5eUIl9dOnAj0k+ICA9Kdu3eH77zySvjW974X52++/HKI80svvb986aXwLfYX9nXTpnlIo/Tp/jT/yy+vl1c8rm0tO78Xz035O2Wz70+///147jwMqAt1+sGRI3FQjb9WIeyLugukskiHBajSpSBV+FNxsGk8moDP0grkK9BIkFKtXhYpg00Hl+fjXwFXcUNJyyyLlFjS+L9NnRAo+g7l3v/l/v1h52uvRasUmAKmFKjAU9ssBc50f7qPdaXv7u8AWvuVt5u2AFUdV/qYrpNG+zhP5m//4Afx3L/72muBusitT/tH+R5r2j+KtS5ADrNceuNBvAbTy/PxY9spSB2Qn98oXcL4KdBIkGKR9gt/mlleiI34/P21So0fQADSdMDp+JUr0b0/cPZseO348a5VSl8pbvELu3eHP3v11QgnAMUsWKXg0rqWSqdtltqnZbpP6YrH0u1iGo7pOEugz7liiQJR6hCt0cOHozVKHenKoEuDgSbcelmjVUG6sLYar8Gh5YUYSVF07d1HOn4g8BnnKdBokPYatZ+9sA7SuTu3KoM09pN23Pv4htPKSvwSEiPZWGz0lX6/4+LTX/oXe/dGoGKh4u4DqT9j7gBLS/YXj3X3vfpqTK+0LHvN/dLHcjt5lCZdcj5sc36cJ+48luhL09OxLruPH++O1sutp+4RpJ137KtY9zyQzt69HUHKtWDU3iDNa4TOPf4KNBakcu2JI1UfKY325NI6SE/cWKkMUoCx1LFK4+j99evRQmMAhrecgCmDMvSX4hK/ODUVoYR1B1SZARUzFmu61Dr702Pa1jLNX9ynMtJlcV35tdR5cY7An3Pm3OkX3XX0aIwd5dN5GmTq5dZXBenxG1cjSE8sGaTjjwDXoA4FGgtSxZEygJEG5J85vxgb8czl86OBtBOYz6CT+kqPXr4cQQNwCBFi8IlwIQGV0CjiTJkjrFhOTUVosZ2uA7LizHFBjmPa1jLu60Awlpesqyz9Rnc7+d14blighw7Fc6a/lzrwYODDLLj0hDwBUepM3XHrR3mjCYt05vJSvAZz5xfClStXYmhaGkdq176OpukyxkmBRoIUAWmMcu1Ti3RhYSEcvLA+4DS/ercSTKNF2gmDAiQawectH0CD1cbbToQJCai4+1ipBLRj5THj/jMD2u56Z7/2sezOSR6lV/4N5XXKj/kKeWK6DijJy5tKOh/ODXhyrgCUV0H3nz0brWzqxIOCOsqlTz/mXNUaVf/o9IX5sLi4GEGKa881euutt2LImkE6TgjwudahQCNBqvCnNCCfxsqoPSCdXTgbLaJj169UBingYNbAEy4+oUCM4jMQQz/iIYC6uBgIFyLOFMuOGUDtOXkyDt4ALM2vnzgR4syxZH/3eGe/0mg/S6WPx5K8fGAkHtfvHT++YVtlkY7z4vwYUOIhwLljhfIlfOrUD6JYl1VBiuaEPh1ZPNcF6a1bt6LXoPAng7SOpukyxkmBRoNUFuna2loc0KCPFJCePjMXGzMN+vyD+0PDVBBlqQB9WaZYa7zxRIwl1inv4hNnCpB4AwpLlZm+VGa6AQBtnOfnowvNPlxpZq3reN9tlZN4/SEFAAAO6UlEQVTkUX6Vr+0NZXMe8/Pr57W4GM+Rc8UC5dypAwBlhL6XJUr9q0IUrRU/evLMXDh//ny0SAEpFikg5ZrRLePJCrRJgbECKf1xuJOnT58OM+fWYXrk6sWhQYoFphmQFGEq6xT4AKGTKyvRogNKuMcA6iiAvXgxWntaYvlpZh8AjkvWtc0+7V9eXk9/6VI3X8zfKVt52af8/LZ+o7vOPs7p8uU4Y33yIIgAvXEjWtrEi9KFsUCfaPIWU1WIopus0UPzc2Fubh2kKysrAZDysNNn9PxmU5sQ4rqiQGNBinuPdUO/G42UxgpIsYJoxEePHw8HL5yLFtLc7ZtdQAqUg5aCyAbLdG0t0HcIdIAPVhwuPzNgZaAG91+A7bXkuEBWPE7+uK8DaNZ7lqnf0HJl5f3fJU9nf3fZOTfOca4DT1mgMcRpdbU7sKTBJeotDQbplB47fftG1Prg0tlw9MSJcObMmfj3u4D09u3b8RpxrWSRcv08WYG2KNB4kPKmzIMHD+Lf/dJoL1y4EM6ePRtOnDgRZuZOxsY9fXExLK6tVoJpCglZp7xCiqsvoAJVYMQsuGoJrPrNxTTpNmAmH/vS/SpL+1gqrdJryX6l0z62da6cN/WgPrJC5cpXBSg64dIfvLgeLXHo1Ilw8uTJeA24FteuXQt37twJ9+/fD1wrg7Qt6HA9UwUaDVL+jVIgpbHSaPlwyblz5yJI+ffBgx2YEpJTpb80BSnrsk7l7gMggVVwxT0Gsuk8X9jWsW6ePsf75YtueCfPhjJ6/Hb6W6SN59xZYn0yC6DF+g67jaYKd0JrNAek8/Pz8VrwX1r89a5AyjVzH2naxLzeBgUaDVJGf/UFKAYz0pF7+kmPHj0aZmZmwtT8en9pLkyBi8DDMp0FJmDVnTuw6h4j2J/jxf3aV1ySTsBOlml+zkFp2E/58aUClVX4LZ2/3PeiBVrcHgRUQp0E0amFuTA9PR1BivYM+vV6q0l/xWzXvg34cB2lQKNBqqB8BjHSftKlpaXYT3r8+PFw+PDhMDV9MBxYWIcpbn7VPtNBMEnBk8KpbF0QTtP1glx6nHWl4ZzSY9qvNOm20qV5BtVpmGP0icqdnzp3OkxNT0et6VJR/+jVq1dj3zUPueJAk0GqJuZlGxRoJEgRnoYISLFwGMTAdZR7v7y8HF1LXExcTazSAwcOhAOnjnXDcxhhznH1h4HNVqYBlv1+b9Cxfnn67UczIiEU5rT/5LGwf//+Ddao3Pq0f9QDTW3AhevYT4FGgxSYyr3H4uE1RD7ZhktJGBSj91ilhw4dClNTU2Hfvn1h/+yhcGDxTBcEADVnIKofcCZtP2+JKbwJiB5YOhf2HzsSNUVbNEZrhT0V3Xr6stU/amu0X3Pz/klVoLEgRXCBlJFgufdYpYzeY5Uy6IRVqr5SrNK9e/eGvQf2h/2nj3dhChjo6+NDJ+fu3YlgrdOKGzeoUvfzD9bCmTu3oiYzl9bfnZcVuu/EbNi7b2/UEk2x+LH8T506FftG0V5hT3gKWKP0ZfuNpknFhOtVpkDjQSr3XqP39MdhlSqmlP46+u1o6AyGYD0B0z179oQ9+/eF/aeObbBQBQsv179XIB2wQPceOxJe37MnvP766xsgeuTIkfjAUt8o2uv9et5mcthTWTPz8UlXYCxAiqWTBuerr5RQqNTFp8EDU1mmAIH5tddeC69PHYig2Dd3IhyYPx1dV0GkbUugeWBxLqDF3tnDUZvdu3ev65RA9ODBgwFN5dKjNZr36hvlGvHQw4uwaz/p2HD9igo0GqScLI1TVmkaCsWbTowa07AZ/FB/KaP4gil9plingilA1Qw4du3a1Z3ZTmeOaTtNl66nx4vraX6ODcqntMVlWqaOaV+6VNnF36GuxXTpNuvxIdN54KAVA0s8iNBQECXcSRDFpU/frZdLb4gWm5a326RA40GKdUMjTa1S3nTSN0pxM+mzI64RmOLmA4AYXzo1FaEAUHH1UwtVQO21FGAEHQEpXSqf0hSXKbx7lVfMr20tlUfbWup32Na60mq7mDY9rnw6v9gFsmdPHFQCoHSNFCHKgwqNeXDp/5kUgC+Q2gptEzZc16ICjQcpJyx3kVFhDTzRkHmjhr46GngRpgxAMdIMFIADlpYsVOChflSAkkI2XRd4lQYIab1fOh3XUum01H6Wo6wX86Tb/IbmQeWr7rJA0QY3Hq2w6Olv5oHEg4kHlCCK1mjOg0xxo3bpi03K221UYGxAWnTxaciKLdXgk2DKu/iMMNO3BxSAAxYqsACqWF7MAESA1TKGUHWgq33pUuukK5uLadnWPi3T39N6v2O9fi8tU8fT/GmZ7Fe9ZX2iCdrw0EErNEM7BpZw5wVRNAai6Vee0nAnW6RtxIfrLAXGBqSyStVfykgxI8ayTGno6jPlzSf1mwqoWKi4/IIq8MACK86yzNgv+LLO/nRO92mdZXFOyyumU3m98uiYlkqTlpeuF4/rWHFJOoETPZjRBoASSoYVSlgZX9mi/xlNsUTpSkFrfbwZiGKN6rrohvLSCrRRgbEAqS4MjTbtLwWmskxp6AyCMKJMsDiWFDAAqFioDJgAClzWY8eOdWcggiUmoLCezuzXdrqufVpyrHi8uE3aXvvSc+h1XPtYal2/O6jMNC3r+h3VX/BEGyxQACorlL5ntERTQRStFepkl153pZdWoKHfIx10YVKYYhUx2EFAOP12xJgSGoV1yleJACpWFRYqgBBUsbo0Y7EWZ4DLzH4te6UpHle+fkuVoTLTpfIU0+g3ise1v99SZRfLA5rM1F/wpB+Uhw4PHzQDoFihaImmaZ+oPpPHA83W6KA71cfapMBYWaS6MGrAcvMFU7n6WFBAABgABawr/u8JqPINTaABWAGIlkCWbZbpnO5L15WmuI/tXvuw9tI8vdLoOMvi8fSY1vulYX+vY+yjvtSfhwtaoAna4MKnAEVD+kPRVG8uqU/UENWd6KUVWFdgrEEKUHExmbGUACruJxYUEBBQ+YK7oEocJNAArlhfQGTQLNAIxL3SpmmKZbKtfWm6XuX021fMV9wu5kuP67d1HlqiAVoATyx4XHgePgIoGgJQufLFPlG092QFrMC6AmMJ0vTiyToFpqmrD1BloeKeAghGnYEq0BBcWQeywERLugW0zv6qM/nTPJSV7uu33us3i/vKttPfLa6TVzP1B5xogjZopMEktOOhxCyA2gpN7zqvW4GNCow9SKkOMFVDZ4l1KgsVq0pQxcoCFsxYrMAjnQGKZsFXkGFb0FEaHUu3+62n5fVL029/r7y99pG/uD/dZp2ZumtW/ycaYX2mFqggqofVxlvHW1bACkiBiQEpFVKDl7svlx+oChKAVXAFHoIs1ms66xhLYJMeS9eVX8v02KD1Ynr9hvZrWSxD+wXAfscH7acM6cBS8EQn4MnMA0kPJ3T1ZAWsQH8FJgKkVE+NXTBlCQjk8msJJGSt4rpqHZjIndWy1z6lJ42Oa5+2lb/XkjRpunS9V/r0d/od136V1es3dI6kZV2zrE70ETxTgErX/reQj1gBKzAxIO11KQWBIlwFWOChOd0nq0yQ0VIwFnCUN92vdZWRptF6r/J0TPnTbe1Ly2Rd56zydFzbxTKK26pHqo/We+npfVbACvRWYKJBWqyyIFFcCigsOVZcKj3laV3plJZjWtcxpU339zqWHi+WozKK+dLtNH+6PiivjqV1Yr04kc6TFbACgxVoFUgHS7E5R4sgKm73+9Vh0/XLX9xfd3nF8r1tBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAgZpLTK6ECtgBdqsgEHa5qvvulsBK1CLAoD0bi0luRArYAWsQDsVuAtIL7Sz7q61FbACVqAWBWYB6c5ainIhVsAKWIF2KrATkO5oZ91daytgBaxALQo8DUifrKUoF2IFrIAVaKcCTz7G5AGndl5919oKWIFsBZYiRDsgfTq7OBdgBayAFWifAs+lIH28ffV3ja2AFbAC2Qo80QVpxyp9NrtIF2AFrIAVaI8C71ujomkIAavUwfntuQlcUytgBUZXYCmEsNEaTWDqUKjRhXVOK2AF2qPAU+Jmz+WPLVO7+O25GVxTK2AFqivwbE94FneGEHZVL9s5rIAVsAITr8BskZd9tzv9pbMTL4kraAWsgBUYXoHdsLEvOPsdsJs/vMJOaQWswEQrMJw7PwCmT4UQGKHyZAWsgBVomwJEMu3ox8dK+xnm//H8fNsUdH2tgBVotQIMvFd35cvomgDVFmqr7y9X3gpMrAJYoLwyXz9AewG289UofpDvmR51MP/E3liumBWYVAWAJh+2h2HE0K9/xakX8Er2/X9XFMGC497oGAAAAABJRU5ErkJggg==');
        background-repeat: repeat;
        background-size: 350px 200px;
        transform-origin: center;
        opacity: 0.08;
        pointer-events: none;
        z-index: 9999;
        transform: rotate(-45deg);
    ">
    </div>
    """

    style_ad="""
    table {
        page-break-inside: avoid;
    }

    /* Alternative modern property */
    @media print {
        table {
            break-inside: avoid;
        }
    }
    """
    
    if '</style>' in html_content:
        html_content = html_content.replace('</style>', style_ad + '</style>')

    if '</body>' in html_content:
        html_content = html_content.replace('</body>', watermark_div + '</body>')
    else:
        html_content += watermark_div
    with open(output_file_path, 'w', encoding='utf-8') as file:
        file.write(html_content)

    return output_file_path

def render_cover_page(company_name, report_date, prepared_by, output_file_path):
    with open("/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/html_report/templates/cover.html", "r", encoding="utf-8") as file:
        template = Template(file.read())
    
    rendered_content = template.render(
        company_name=company_name,
        report_date=report_date,
        prepared_by=prepared_by
    )
    with open(output_file_path, 'w', encoding='utf-8') as file:
        file.write(rendered_content)
    
    return output_file_path

def merge_pdfs(pdf_paths, output_path):
    merger = PdfMerger()
    for pdf in pdf_paths:
        merger.append(pdf)
    merger.write(output_path)
    merger.close()

async def conversion(ticker_name):
    base_dir = '/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/html_report'
    watermark_html_path = os.path.join(base_dir, 'watermark_download.html')
    cover_html_path = os.path.join(base_dir, 'templates/cover_filled.html')
    report_pdf_path = os.path.join(base_dir, 'Final_report_watermark.pdf')
    cover_pdf_path = os.path.join(base_dir, 'cover_filled.pdf')
    final_pdf_path = os.path.join(base_dir, 'Final_Combined_Report.pdf')

    report_html_path = os.path.join(base_dir, 'Infosys_Unified_Risk_Report.html')
    insert_watermark_div_in_html(report_html_path, watermark_html_path)

    render_cover_page(ticker_name, date.today(), 'ALEC', cover_html_path)
    insert_watermark_div_in_html(cover_html_path, cover_html_path)
    a = await convert_html_to_pdf_async(watermark_html_path, report_pdf_path)
    b = await convert_html_to_pdf_async(cover_html_path, cover_pdf_path)
    print(f"Conversion results - Report: {a}, Cover: {b}")

    if a and b:
        merge_pdfs([cover_pdf_path, report_pdf_path], final_pdf_path)
        print(f"Merged PDF created at: {final_pdf_path}")

    return final_pdf_path





def create_stock_analysis_table(data):
    html = "<h3>FINANCIAL AND MARKET INFORMATION FOR THE COUNTER PARTY</h3>\n"
    html += "<table border='1'>\n"
    html += "  <tr><th>Category</th><th>Metric</th><th>Value</th></tr>\n"

    # Metadata
    metadata = data.get("metadata", {})
    meta_rows = []
    if "ticker_symbol" in metadata:
        meta_rows.append(f"  <tr><td rowspan='{{rowspan}}'>Metadata</td><td>Ticker Symbol</td><td>{metadata['ticker_symbol']}</td></tr>\n")
    if "organization" in metadata:
        meta_rows.append(f"  <tr><td>Organization</td><td>{metadata['organization']}</td></tr>\n")
    if "evaluation_date" in metadata:
        meta_rows.append(f"  <tr><td>Evaluation Date</td><td>{metadata['evaluation_date'].split('T')[0]}</td></tr>\n")
    if "data_providers" in metadata:
        providers = ", ".join(metadata.get("data_providers", []))
        meta_rows.append(f"  <tr><td>Data Providers</td><td>{providers}</td></tr>\n")
    if meta_rows:
        html += "".join(meta_rows).replace("{{rowspan}}", str(len(meta_rows)))


    financial = data.get("financial_indicators", {})
    fin_rows = []
    if "valuation_ratio" in financial:
        fin_rows.append(f"  <tr><td rowspan='{{rowspan}}'>Financial</td><td>Valuation Ratio (P/E)</td><td>{financial['valuation_ratio']}</td></tr>\n")
    if "net_margin" in financial:
        fin_rows.append(f"  <tr><td>Net Margin</td><td>{financial['net_margin']}</td></tr>\n")
    if "liquidity_ratio" in financial:
        fin_rows.append(f"  <tr><td>Liquidity Ratio</td><td>{financial['liquidity_ratio']}</td></tr>\n")
    if "leverage_ratio" in financial:
        fin_rows.append(f"  <tr><td>Leverage Ratio</td><td>{financial['leverage_ratio']}</td></tr>\n")
    if "asset_returns" in financial:
        fin_rows.append(f"  <tr><td>Asset Returns (ROA)</td><td>{financial['asset_returns']}</td></tr>\n")
    if "equity_returns" in financial:
        fin_rows.append(f"  <tr><td>Equity Returns (ROE)</td><td>{financial['equity_returns']}</td></tr>\n")
    if "company_valuation" in financial:
        fin_rows.append(f"  <tr><td>Company Valuation (Market Cap)</td><td>{financial['company_valuation']}</td></tr>\n")
    if "annual_peak" in financial:
        fin_rows.append(f"  <tr><td>Annual Peak Price</td><td>{financial['annual_peak']}</td></tr>\n")
    if "annual_trough" in financial:
        fin_rows.append(f"  <tr><td>Annual Trough Price</td><td>{financial['annual_trough']}</td></tr>\n")
    if fin_rows:
        html += "".join(fin_rows).replace("{{rowspan}}", str(len(fin_rows)))

    market = data.get("market_information", {})
    market_rows = []
    if "current_price" in market:
        market_rows.append(f"  <tr><td rowspan='{{rowspan}}'>Market Information</td><td>Current Price</td><td>{market['current_price']}</td></tr>\n")
    if "medium_term_average" in market:
        market_rows.append(f"  <tr><td>Medium-Term Average Price</td><td>{market['medium_term_average']}</td></tr>\n")
    if "long_term_average" in market:
        market_rows.append(f"  <tr><td>Long Term Average Price</td><td>{market['long_term_average']}</td></tr>\n")
    if "price_volatility" in market:
        market_rows.append(f"  <tr><td>Price Volatility</td><td>{market['price_volatility']}</td></tr>\n")
    if market_rows:
        html += "".join(market_rows).replace("{{rowspan}}", str(len(market_rows)))

    html += "</table>"
    return html


def main2(ticker_symbol,ticker_name,corporation,jurisdiction,period,region):
    ticker_symbol = ticker_symbol 
    ticker_name = ticker_name
    corporation =corporation
    jurisdiction = jurisdiction  
    period = period  
    region = region 

    Financial_analysis_Agent(ticker_symbol,ticker_name)
    Legal_analyisi_Agent(corporation,jurisdiction,period)
    company_Background_anslysis_Agent(ticker_name,jurisdiction)
    json_file="/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Legal_analyses_output/legal_analysis.json"

    with open(json_file, "r") as f:
        json_data = json.load(f)

    legal_cases_table = create_legal_cases_table(json_data)

    case_categories="/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Legal_analyses_output/case_categories.png"
    case_severity_distribution="/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Legal_analyses_output/case_severity_distribution.png"
    case_timeline="/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Legal_analyses_output/case_timeline.png"
    risk_matrix="/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Legal_analyses_output/risk_matrix.png"
    financial_health_dashboard="/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Stock_analyses_output/financial_health_dashboard.png"
    price_analysis="/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Stock_analyses_output/price_analysis.png"
    risk_analysis="/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Stock_analyses_output/risk_analysis.png"

    stock_file="/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Stock_analyses_output/stock_analysis.json"
    with open(stock_file, "r") as f:
        stock_data = json.load(f)
    stock_analysis_table = create_stock_analysis_table(stock_data)


    FILES = {
        "financial": "/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Stock_analyses_output/company_summary_report.json",
        "corporate": "/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Company_Background_analysis_output/company_background_analysis.json", 
        "legal": "/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Legal_analyses_output/legal_analysis.json"
    }
    
    try:
        analyses = analyze_all_files(FILES)
        

        print("\nGenerating executive summary...")
        summary = generate_executive_summary(analyses)
        

        full_report = f"""# {ticker_name} Unified Risk Report
**Report Generated**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## Executive Summary
{summary}

## Detailed Analysis

### FINANCIAL RISK ASSESSMENT
{analyses.get('financial', 'No data')}

{image_markdown_if_exists(financial_health_dashboard, "Financial Health Dashboard")}

### CORPORATE STRUCTURE ANALYSIS  
{analyses.get('corporate', 'No data')}

{image_markdown_if_exists(price_analysis, "Price Analysis")}
{image_markdown_if_exists(risk_analysis, "Risk Analysis")}

{stock_analysis_table}

### LEGAL RISK EVALUATION
{analyses.get('legal', 'No data')}

{image_markdown_if_exists(case_categories, "Case Categories")}
{image_markdown_if_exists(case_severity_distribution, "Case Severity Distribution")}
{image_markdown_if_exists(case_timeline, "Case Timeline")}
{image_markdown_if_exists(risk_matrix, "Risk Matrix")}

### LEGAL CASE DETAILS

{legal_cases_table}

"""

        report_file = save_report(full_report)
        print(f"\n Unified report generation complete!")
        print(f" Report saved to: {os.path.abspath(report_file)}")
        report_path=os.path.abspath(report_file)
        # Print summary
        print("\n=== EXECUTIVE SUMMARY ===\n")
        print(summary[:2000])
        folder = "/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/"
        all_images = get_all_image_paths(folder)

        for img_path in all_images:
            print(img_path)

        return report_path,all_images
        
    except Exception as e:
        print(f" Error during unified analysis: {str(e)}")
        return None, None





@app.get("/download-report")
async def download_report(file_path: str):
    try:
        path_to_upload=await conversion(corporation)
        file_id = await upload_file_to_db(path_to_upload)
        print("Uploaded file ID:", file_id)
        return FileResponse(path=path_to_upload, filename="report.pdf", media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=404, detail="File not found")




@app.post("/CpyDelData")
async def submit_company_details(request: Request):
    risk=RiskAnalyzer()
    risk.unload_model()
    try:
        company_details = await request.json()
        required_fields = ["ticker_symbol", "ticker_name", "corporation", "jurisdiction", "period", "region"]
        missing_fields = [field for field in required_fields if field not in company_details]

        if missing_fields:
            raise HTTPException(status_code=400, detail=f"Missing fields: {', '.join(missing_fields)}")
        ticker_symbol = company_details["ticker_symbol"]
        ticker_name = company_details["ticker_name"]
        global corporation
        corporation = company_details["corporation"]
        jurisdiction = company_details["jurisdiction"]
        period = company_details["period"]
        region = company_details["region"]

        print(company_details)
        result = main2(ticker_symbol, ticker_name, corporation, jurisdiction, period, region)
        if result is not None:
            a, c = result
        else:
            a, c = None, None
        print(a,c)
        report_path=generate_report_html(a)
        image_paths = c
        json_file="/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/pro/Legal_analyses_output/legal_analysis.json"

        with open(json_file, "r") as f:
            json_data = json.load(f)

        return JSONResponse(
             content={
        "message": "Company details submitted successfully!",
        "report_path": report_path,
        "image_paths": image_paths,
        "json_file": json_data,
    },
    status_code=200
        )
    
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))



api1  = "GROQ API KEY"
api2  = "GROQ API KEY"

def asker_agent(requirement_dict, file_docu_name):
    client = Groq(api_key=api2)
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": f"""
                You are an intelligent requirements extraction assistant. Your task is to generate the **most relevant and concise** questions to gather all necessary details for a contract.

                **Guidelines:**  
                - **Minimize the number of questions** while ensuring completeness.  
                - **Avoid redundant or vague questions.**  
                - **Ensure each question is direct, precise, and necessary.**  
                - Format the response as a structured list of questions **without explanations or additional text.**  

                **Contract Name:**  
                {file_docu_name}  

                **Requirements:**  
                {requirement_dict}  

                **Output Format:**  
                - Provide the questions in bullet points or numbered format.  
                - Do not include extra commentary or formatting beyond the questions.  
                """
            }
        ],
        model="llama-3.3-70b-versatile",
    )
    return chat_completion.choices[0].message.content.strip()

def rechat_agent(requirement_dict, file_docu_name):
    client = Groq(api_key=api2)
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": f"""
                You are an intelligent requirements extraction assistant. Your task is to **ask only the questions explicitly mentioned** in the provided input. Do not generate additional or unnecessary questions.

                **Guidelines:**  
                - **Ask only the given questions** without adding or modifying them.  
                - **Do not generate new questions** beyond what is provided.  
                - **Maintain the original wording and sequence** of the questions.  
                - Format the response as a structured list of questions **without explanations or additional text.**  

                **Contract Name:**  
                {file_docu_name}  

                **Questions to Ask:**  
                {requirement_dict}  

                **Output Format:**  
                - Provide the questions in bullet points or numbered format.  
                - Do not include extra commentary or formatting beyond the questions.  
                """
            }
        ],
        model="llama-3.3-70b-versatile",
    )
    return chat_completion.choices[0].message.content.strip()

def validator_agent(a, text):
    client = Groq(api_key=api2)
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": f"""You are a validator. Your task is to check whether the provided text contains answers to all given questions.  

                - If a question is fully answered, do not include it in the response.  
                - If any question is missing its answer, return only that specific question inside a **single** <question></question> tag.  
                - If at least one answer is missing, return <missing>yes</missing>. Otherwise, return <missing>no</missing>.  
                - Ensure there is exactly **one** <question> tag and **one** <missing> tag in the response.  

                Answer Text: {text}   
                
                Questions:  
                {a}  
 

                Return only the formatted response without explanations or additional comments.
                """
            }
        ],
        model="llama-3.3-70b-versatile",
    )
    return chat_completion.choices[0].message.content.strip()

def format_legal_text(text):
    client = Groq(api_key="GROQ API KEY")
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": f"""
                Format the given text using appropriate HTML tags while preserving its original content. Apply suitable tags for headings, paragraphs, lists, and emphasis where needed. Do not alter the wording or structure of the text.  
                
                Text:  
                {text}  
                
                Provide the formatted text without explanations or additional comments.
                """
            }
        ],
        model="llama-3.3-70b-versatile",
    )
    return chat_completion.choices[0].message.content.strip()

def modifier_model1(text, requirement):
    client = Groq(api_key=api1)
    chat_completion = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {
                "role": "user",
                "content": f"""i am having text, in that text i want to fill details using the requirement given below\n
                 your job is to fill the details in the text below and resticted to generate extra content.\n\
                 exclude explanation.
                 text:\n {text}\n.
                 requirement :\n {requirement}\n
                """
                
            }
        ],
        
    )
    return chat_completion.choices[0].message.content.strip()

apiaiconracter = "GROK API KEY"

def type_checker_agent(text_input: str) -> str:

    client = Groq(api_key=apiaiconracter)
    
    prompt = f"""You are a legal document analysis expert. Perform the following tasks with precision:

                            1. DOCUMENT ANALYSIS:
                            - Carefully read the provided text
                            - Identify all potential contract/agreement references
                            - Determine the primary document being discussed

                            2. NAME EXTRACTION RULES:
                            - Extract only the official document name
                            - Exclude version numbers, dates, and ancillary text
                            - Prefer formal names over colloquial references
                            - If multiple names exist, choose the most specific one
                            - If no clear name exists, return "Unspecified Agreement"

                            3. OUTPUT FORMAT:
                            - Return only the exact name without explanations
                            - Capitalize proper nouns appropriately
                            - Remove quotation marks unless part of the official name

                            TEXT TO ANALYZE:
                            {text_input}

                            Final Answer:"""
    
    try:
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0.2,  
            max_tokens=50
        )
        

        result = response.choices[0].message.content.strip()
        

        result = result.replace('"', '').replace("'", "").replace("\n", " ").strip()

        if not result or result.lower() in ["none", "n/a", "unknown"]:
            return "Unspecified Agreement"
            
        return result
        
    except Exception as e:
        print(f"Error in type_checker_agent: {e}")
        return "Identification Failed"


def clause_generation_agent(context_text: str, contract_name: str) -> list:

    client = Groq(api_key=apiaiconracter)
    
    structured_prompt = f"""As a legal document architect, generate the perfect clause structure for a {contract_name} following these strict rules:

                            1. MANDATORY STRUCTURE:
                            [Opening Clauses]
                            1. Parties
                            2. Recitals
                            3. Definitions
                            4. Term

                            [Core Clauses]
                            5. {{context-specific-clauses}}

                            [Closing Clauses]
                            N-3. Governing Law
                            N-2. Dispute Resolution
                            N-1. Miscellaneous
                            N. Signatures

                            2. GENERATION RULES:
                            - Extract ALL unique clauses from this context: {context_text[:1500]}...
                            - Never repeat any clause
                            - Maintain EXACTLY this order
                            - Include only essential clauses (no fluff)
                            - Use standardized legal names
                            - Return ONLY a comma-separated list

                            3. OUTPUT FORMAT:
                            Parties, Recitals, Definitions, Term, [core clauses...], Governing Law, Dispute Resolution, Miscellaneous, Signatures

                            Generated clauses:"""

    try:
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": structured_prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0.3,  
            max_tokens=400
        )
        
        raw_clauses = response.choices[0].message.content.strip()
        

        mandatory_opening = ["Parties", "Recitals", "Definitions", "Term"]
        mandatory_closing = ["Governing Law", "Dispute Resolution", "Miscellaneous", "Signatures"]
        

        core_clauses = [
            clause.strip() 
            for clause in raw_clauses.split(",") 
            if clause.strip() and clause.strip() not in mandatory_opening + mandatory_closing
        ]
        seen = set()
        unique_core = [x for x in core_clauses if not (x in seen or seen.add(x))]
        

        final_clauses = mandatory_opening + unique_core + mandatory_closing
        
        return [clause for clause in final_clauses if isinstance(clause, str) and clause.strip()]
        
    except Exception as e:
        print(f"Clause generation error: {e}")
        return ["Parties", "Recitals", "Definitions", "Term", 
                "Governing Law", "Dispute Resolution", 
                "Miscellaneous", "Signatures"]


def asker_agent1(clause_list: list, context_text: str) -> dict:
    client = Groq(api_key=apiaiconracter)
    
    enhanced_prompt = f"""As a legal research assistant, generate ONLY the most essential questions 
                                needed to complete this contract, categorized by information source:

                            1. CONTEXT (Already Known - DO NOT Ask About):
                            {context_text[:2000]}{'...' if len(context_text) > 2000 else ''}

                            2. CLAUSES REQUIRING INFORMATION:
                            {clause_list}

                            3. CATEGORIZATION RULES:
                            - **User Questions**: 
                            * Facts only the contracting parties know
                            * Specific business terms
                            * Contact/identity information
                            * Example: "What is the exact payment amount for Phase 1?"

                            - **Web Questions**:
                            * Standard legal provisions
                            * Jurisdiction-specific requirements  
                            * Template language requests
                            * Example: "What are the standard warranty clauses for SaaS in California?"

                            4. OUTPUT FORMAT (STRICT JSON):
                            {{
                            "user": ["question1", "question2"],
                            "web": ["question1", "question2"]
                            }}

                            5. GENERATION RULES:
                            - Absolute minimum questions needed
                            - No duplicates
                            - No questions answerable from context
                            - Legal-precise phrasing

                            Generate questions:"""

    try:
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": enhanced_prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0.3,
            response_format={"type": "json_object"}
        )
        
        questions = json.loads(response.choices[0].message.content)
        return {
            'user': [q for q in questions.get('user', []) if isinstance(q, str)],
            'web': [q for q in questions.get('web', []) if isinstance(q, str)]
        }
        
    except Exception as e:
        print(f"Question generation error: {e}")
        return {'user': [], 'web': []}


serpapiaicontracter = "SERP API KEY"

import requests
import json
from groq import Groq

def research_contract_questions(web_questions: list) -> str:
    groq_client = Groq(api_key=apiaiconracter)
    all_search_results = []
    
    for question in web_questions:
        try:
            params = {
                'q': f"{question} site:gov OR site:edu OR site:law",
                'api_key': serpapiaicontracter,
                'num': 3
            }
            response = requests.get('https://serpapi.com/search', params=params)
            response.raise_for_status()
            
            results = [
                f"Source: {result.get('source', 'N/A')}\n"
                f"Title: {result.get('title', '')}\n"
                f"Content: {result.get('snippet', '')}"
                for result in response.json().get('organic_results', [])[:3]
            ]
            all_search_results.append(f"QUESTION: {question}\nSEARCH RESULTS:\n" + "\n\n".join(results))
        except:
            all_search_results.append(f"QUESTION: {question}\nSEARCH RESULTS: Research failed")

    research_content = "\n\n".join(all_search_results)
    
    refinement_prompt = f"""As a legal expert, synthesize this research into a concise contract summary:

                            Research Content:
                            {research_content[:6000]}

                            Instructions:
                            - Combine related information
                            - Use formal legal language
                            - Cite authoritative sources
                            - Structure by legal topics
                            - Keep under 400 words

                            Output Format:
                            Final summary of all the answer.
                            Provide the formatted text without explanations or additional comments."""

    try:
        response = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": refinement_prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0.3,
            max_tokens=800
        )
        return response.choices[0].message.content.strip()
    except:
        return "Legal research summary unavailable"

from sentence_transformers import SentenceTransformer
modelaicontracter = SentenceTransformer('all-MiniLM-L6-v2')

from pinecone import Pinecone
from groq import Groq
from typing import List, Dict
import json

class ContractDrafter:
    def __init__(self, pinecone_api_key: str, groq_api_key: str, index_name: str = "unisys"):
       
        self.pc = Pinecone(api_key=pinecone_api_key)
        self.index = self.pc.Index(index_name)
        
        self.groq_client = Groq(api_key=groq_api_key)
    
    def draft_full_contract(self, contract_name: str, clause_order: List[str], requirements: str, 
                          party_a_favorability: int, party_b_favorability: int) -> str:

        if party_a_favorability + party_b_favorability != 100:
            raise ValueError("Favorability percentages must sum to 100")

        contract_sections = []
        
        for clause_name in clause_order:
            similar_clauses = self.retrieve_clauses(
                contract_name=contract_name,
                clause_name=clause_name,
                top_k=5
            )
            
    
            final_clause = self.draft_balanced_clause(
                clause_name=clause_name,
                example_clauses=similar_clauses,
                requirements=requirements,
                party_a_favorability=party_a_favorability,
                party_b_favorability=party_b_favorability
            )
            
            contract_sections.append(final_clause)
        
        return "\n\n".join(contract_sections)
    
    def retrieve_clauses(self, contract_name: str, clause_name: str, top_k: int = 5) -> List[Dict]:
        query_text = f"{contract_name} {clause_name}"
        query_embedding = self.get_embedding(query_text)
        
        results = self.index.query(
            vector=query_embedding,
            top_k=top_k,
            filter={
                "contract_type": {"$eq": contract_name},
                "clause_name": {"$eq": clause_name}
            },

            include_metadata=True
        )
        
        return [
            {
                "text": match.metadata["clause_text"],
                "score": match.score
            }
            for match in results.matches
        ]
    
    def draft_balanced_clause(self, clause_name: str, example_clauses: List[Dict], 
                             requirements: str, party_a_favorability: int, 
                             party_b_favorability: int) -> str:
        examples_formatted = "\n\n".join(
            f"EXAMPLE {i+1} (Similarity: {clause['score']:.2f}):\n{clause['text']}"
            for i, clause in enumerate(example_clauses)
        )
        
        prompt = f"""As a legal contract attorney, draft a {clause_name} clause that is:
                    - {party_a_favorability}% favorable to Party A
                    - {party_b_favorability}% favorable to Party B

                    CLIENT REQUIREMENTS:
                    {requirements}

                    EXAMPLE CLAUSES:
                    {examples_formatted}

                    INSTRUCTIONS:
                    1. Create balanced terms according to specified percentages
                    2. For {party_a_favorability}%/{party_b_favorability}% split:
                    - Party A advantages: {self.get_favorability_guidance(party_a_favorability)}
                    - Party B advantages: {self.get_favorability_guidance(party_b_favorability)}
                    3. Use professional legal language
                    4. Keep comprehensive but concise
                    5. Extract the BEST elements from each example
                    6. Modify to perfectly match the requirements
                    7. provide text for mentioned clause name only
                    8. Provide only the text without explanation or addtional comments.
                    9. provide the text in markdown format.

                    OUTPUT FORMAT:
                    [Clause Name]
                    [Balanced clause text]"""
        
        response = self.groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0.3,
            max_tokens=1000
        )
        
        return response.choices[0].message.content.strip()
    
    def get_favorability_guidance(self, percentage: int) -> str:
        if percentage >= 70:
            return "Strongly favorable (most terms advantageous)"
        elif percentage >= 60:
            return "Moderately favorable (key terms advantageous)"
        elif percentage >= 50:
            return "Neutral (balanced terms)"
        else:
            return "Limited protection (basic rights only)"
    
    def get_embedding(self, text: str) -> List[float]:
        return modelaicontracter.encode(text).tolist()




from pymongo.mongo_client import MongoClient
from pymongo.server_api import ServerApi
uri="MONGO DB URI" # to retrieve the data from the db


global c
c=""
client_sabari = MongoClient(uri, server_api=ServerApi('1'))
try:
    client_sabari.admin.command('ping')
    print("Pinged your deployment. connected to MongoDB!")
except Exception as e:
    print(e)


client1 = MongoClient("MONGODB URI TO STORE FILES USER UPLOADS AND THE CHATS")
db = client1["chat_db"]

db2 =client1["TempalteDB"]


fs = gridfs.GridFS(db2)
messages_collection = db["messages"]
sessions_collection = db["sessions"]


setter=0
sections=""
a=""

global text
text=" "

final_text=" "

template=False
without=0
def get_userid_from_request(request: Request):
    return request.headers.get("userid")

import json
import re

def confirmation_agent(reply):
    client = Groq(api_key=api1)
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": """You are a Confirmation Agent specialized in contract drafting processes. 
                Your sole purpose is to determine if the user wants to start the contract drafting process.
                
                RULES:
                1. Respond ONLY with the exact XML tags specified - no explanations, comments, or additional text
                2. Analyze the user's intent carefully
                3. If the user explicitly or implicitly indicates they want to begin contract drafting, return <process>YES</process>
                4. If the user declines, hesitates, or doesn't want to proceed, return <process>NO</process>
                5. For ambiguous cases where intent isn't clear, default to <process>NO</process>"""
            },
            {
                "role": "user",
                "content": f"""Analyze the following user input and provide your response following the rules exactly:
                
                USER INPUT: {reply}
                
                RESPONSE FORMAT:
                <process>YES</process> OR <process>NO</process>"""
            }
        ],
        model="llama3-70b-8192",
        temperature=0.0,  
        max_tokens=20
    )
    
    response = chat_completion.choices[0].message.content.strip()
    return response

@app.post("/create_session/")
async def create_session(request:Request):
    user_id = get_userid_from_request(request)
    setter=0
    without=0
    template=False

    session_data = {
        "user_id":user_id,
        "created_at": datetime.datetime.utcnow(),
        "messages": []  
    }
    result = sessions_collection.insert_one(session_data)
    return {"session_id": str(result.inserted_id), "message": "New chat session created!"}

def confirmation_agent1(reply):
    client = Groq(api_key="GROK API KEY")
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": """ROLE: Negotiation Process Confirmation Agent

RESPONSIBILITIES:
1. Determine user intent regarding negotiation process
2. Classify response into one of three actions
3. Extract correction text when applicable

RESPONSE FORMATS:
- To START negotiation: <process>YES</process>
- To STOP negotiation: <process>NO</process>
- For CORRECTIONS: 
  <process>CORRECTION</process>
  <correction>[EXACT_TEXT_TO_CORRECT]</correction>

RULES:
1. Respond ONLY with the specified XML tags
2. NEVER add explanations or comments
3. For CORRECTION:
   - Must include both tags
   - <correction> must contain the exact text to modify
4. Default to <process>NO</process> for ambiguous cases
5. For mixed intent (correction + continuation), prioritize CORRECTION"""
            },
            {
                "role": "user",
                "content": f"""USER INPUT: {reply}

ANALYSIS:
1. Does the user want to:
   - START negotiation? → <process>YES</process>
   - STOP negotiation? → <process>NO</process>
   - MAKE CORRECTIONS? → 
     <process>CORRECTION</process>
     <correction>[EXTRACT_TEXT_HERE]</correction>

OUTPUT (ONLY THE REQUIRED TAGS):"""
            }
        ],
        model="llama3-70b-8192",
        temperature=0.1,
        max_tokens=100,
        response_format={"type": "text"}
    )
    
    return chat_completion.choices[0].message.content.strip()

import json
def correction_agent1(chunk, correction):
    client = Groq(api_key="GROK API KEY")
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": f"""i am having text, in that text i want to change details using the modification text given below\n
                 your job is to fill the details in the text below and resticted to generate extra content.\n
                 if nothing needed to change in the text return as it is without any changes
                 exclude explanation.

                 TEXT:\n {chunk}\n.

                 MODIFICATION TEXT :\n {correction}\n
                """
                
            }
        ],
        model="llama-3.3-70b-versatile",
    )

    return chat_completion.choices[0].message.content.strip()




pdf__path = "/teamspace/studios/this_studio/Uvarajan/Whole_Server/uploads/sla.pdf"








def send_negotiation_email(formdata,path):
    groq_api_key = "GROK API KEY"
    sender_email = formdata.get("sender_email", "")
    receiver_email = formdata.get("receiver_email", "")
    application_type = formdata.get("application_type", "")
    counter_party_name = formdata.get("counter_party_name", "")
    counter_party_role = formdata.get("counter_party_role", "")
    counter_party_companyname = formdata.get("counter_party_companyname", "")
    send_party_name = formdata.get("send_party_name", "")
    sender_party_role = formdata.get("sender_party_role", "")
    sender_party_company = formdata.get("sender_party_company", "")
    smtp_server = "smtp.gmail.com"
    smtp_port = 587
    email_password = "YOUR GPASSWORD"  
    try:
        groq_client = Groq(api_key=groq_api_key)

        negotiation_prompt = f"""Compose a professional email for INITIAL DISCUSSION of a {application_type}. 

                                Important:
                                - DO NOT include any specific percentages, numbers, or contract terms
                                - This is just an introductory email to begin discussions

                                Details:
                                Sender: {send_party_name} ({sender_party_role}, {sender_party_company})
                                Recipient: {counter_party_name} ({counter_party_role}, {counter_party_companyname})

                                Format exactly as:
                                SUBJECT: <subject line>
                                BODY: <email content>"""

        response = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": negotiation_prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0
        )

        content = response.choices[0].message.content
        if "SUBJECT:" in content and "BODY:" in content:
            subject = content.split("SUBJECT:")[1].split("BODY:")[0].strip()
            body = content.split("BODY:")[1].strip()
        else:
            subject = application_type
            body = content

 
        msg = MIMEMultipart()
        msg["From"] = sender_email
        msg["To"] = receiver_email
        msg["Subject"] = subject
        msg.attach(MIMEText(body, "plain"))

        
        if os.path.exists(pdf__path):
            with open(pdf__path, "rb") as attachment:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f"attachment; filename= {os.path.basename(pdf__path)}",
            )
            msg.attach(part)

        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(sender_email, email_password)
            server.send_message(msg)
        
        print("Email sent successfully with AI-generated negotiation content")
        return "Email sent successfully with AI-generated negotiation content"

    except Exception as e:
        return e
        print(f"Error: {e}")

@app.post("/send_message/")
async def send_message(request:Request,data: dict):
    import datetime
    global final_text
    user_id = get_userid_from_request(request)
    global setter
    
    session_id = data.get("session_id")
    user_message = data.get("message")
    if user_message=="cloud_services" or user_message=="Service_Level_Agreement" or user_message=="Software_development_agreement" or user_message=="Software_license_agreement" or user_message=="Software_maintenance_agreement":
        global template
        template=True
    if template:
        print("clicked template")
        if not session_id or not user_message:
            return {"error": "Session ID and message are required!"}
        if setter==0:
            setter+=1
            name=user_message
            docu_name = name+".txt" 
            print(docu_name)
            docu_requirement = name+"_requirement.txt"

            file_docu_data = fs.find_one({"filename": docu_name})
            file_docu_requirement_data =fs.find_one({"filename": docu_requirement})


            if file_docu_data:
                file_docu_content = file_docu_data.read()

            if file_docu_requirement_data:
                file_docu_requirement_content = file_docu_requirement_data.read()
            else:
                print("404 not found in DB")
            if isinstance(file_docu_content, bytes):
                file_docu_content = file_docu_content.decode('utf-8',errors="replace")

            global sections

            sections = re.findall(r'<section>(.*?)</section>', file_docu_content, re.DOTALL)

            print(sections)


            decoded_string =file_docu_requirement_content.decode('utf-8')
            cleaned_string =decoded_string.split('=',1)[-1].strip()
            requirement_dict =ast.literal_eval(cleaned_string)
            print(requirement_dict)

            global contract_name

            contract_name = name
            global a
            a= asker_agent(requirement_dict,contract_name)
            
    

            session = sessions_collection.find_one({"_id": ObjectId(session_id), "user_id": user_id})
            user_message_entry = {
                "sender": "user",
                "text": user_message,
                "timestamp": datetime.datetime.utcnow()
            }

            server_response = a 
            server_message_entry = {
                "sender": "server",
                "text": server_response,
                "timestamp": datetime.datetime.utcnow()
            }

            sessions_collection.update_one(
                {"_id": ObjectId(session_id)},
                {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
            )
            return {
                "user_message": user_message,
                "message": server_response
            }
        if setter==1:
            setter+=1
            global final_requirement_content    
            final_requirement_content =""
            text = user_message
            final_requirement_content = final_requirement_content+text
            b=validator_agent(a,text)


            question_match = re.search(r"<question>(.*?)</question>", b, re.DOTALL)
            missing_match = re.search(r"<missing>(.*?)</missing>", b)

            global question
            global missing

            question = question_match.group(1) if question_match else ""
            missing = missing_match.group(1) if missing_match else ""

            print("Question:", question)
            print("Missing:", missing)
            user_message_entry = {
                "sender": "user",
                "text": user_message,
                "timestamp": datetime.datetime.utcnow()
            }

            if question=="None of the above, all answers are present in the text." or question=="None of the above, all questions are answered":
                server_response = "All necessary information to draft the contract is now available. May I proceed to process the document?"
                setter+=2
                server_message_entry = {
                    "sender": "server",
                    "text": server_response,
                    "timestamp": datetime.datetime.utcnow()
                }

                sessions_collection.update_one(
                    {"_id": ObjectId(session_id)},
                    {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
                )
                return {
                    "user_message": user_message,
                    "message": server_response
                }

            else:
                server_response = question
                server_message_entry = {
                    "sender": "server",
                    "text": server_response,
                    "timestamp": datetime.datetime.utcnow()
                }

                sessions_collection.update_one(
                    {"_id": ObjectId(session_id)},
                    {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
                )
                return {
                    "user_message": user_message,
                    "message": server_response
                }
        if setter==2:
                if missing!="no":
                    rechat= rechat_agent(question,contract_name)
                    print(rechat)
                    input_text = user_message
                    revalid =validator_agent(rechat,input_text)
                    final_requirement_content = final_requirement_content+input_text
                    print(revalid)
                    question_match = re.search(r"<question>(.*?)</question>", revalid, re.DOTALL)
                    missing_match = re.search(r"<missing>(.*?)</missing>", revalid)
                    question = question_match.group(1) if question_match else ""
                    missing = missing_match.group(1) if missing_match else ""

                    print("Question:", question)
                    print("Missing:", missing)

                    user_message_entry = {
                        "sender": "user",
                        "text": user_message,
                        "timestamp": datetime.datetime.utcnow()
                    }

                    server_response = question  
                    server_message_entry = {
                        "sender": "server",
                        "text": server_response,
                        "timestamp": datetime.datetime.utcnow()
                    }

                    sessions_collection.update_one(
                        {"_id": ObjectId(session_id)},
                        {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
                    )

                    return {
                        "user_message": user_message,
                        "message": server_response
                    }
                else:
                    setter+=1
                    user_message_entry = {
                        "sender": "user",
                        "text": user_message,
                        "timestamp": datetime.datetime.utcnow()
                    }

                    server_response = "All necessary information to draft the contract is now available. May I proceed to process the document?"  # Example: Simple echo response

                    server_message_entry = {
                        "sender": "server",
                        "text": server_response,
                        "timestamp": datetime.datetime.utcnow()
                    }

                    sessions_collection.update_one(
                        {"_id": ObjectId(session_id)},
                        {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
                    )

                    return {
                        "user_message": user_message,
                        "message": server_response
                    }
        
        if setter==3:
            intent = user_message
            confir_text = confirmation_agent(intent)
            match = re.search(r"<process>(.*?)</process>", confir_text)
            print("entered into setter 3")
            if match:
                process= match.group(1)
                print(process)
                if process:
                    docu=len(sections)
                    count=0
                    collector =""
                    formatter = "f"
                    while(count!=docu):
                        temp_collector = modifier_model1(sections[count] ,final_requirement_content)
                        collector += format_legal_text(temp_collector)
                        count+=1
                    with open(CONTRACT_FILE_PATH, "w", encoding="utf-8") as file:
                        file.write(collector)
                    print("HTML file saved as contract.html")
                    user_message_entry = {
                        "sender": "user",
                        "text": user_message,
                        "timestamp": datetime.datetime.utcnow()
                    }

                   
                    server_response = CONTRACT_FILE_PATH

                    server_message_entry = {
                        "sender": "server",
                        "text": server_response,
                        "timestamp": datetime.datetime.utcnow()
                    }

                    sessions_collection.update_one(
                        {"_id": ObjectId(session_id)},
                        {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
                    )
                    setter=0
                    return {
                        "user_message": user_message,
                        "message": "this is file",
                        "file":server_response
                    }
                else:
                    user_message_entry = {
                        "sender": "user",
                        "text": user_message,
                        "timestamp": datetime.datetime.utcnow()
                    }

                    server_response = CONTRACT_FILE_PATH 

                    server_message_entry = {
                        "sender": "server",
                        "text": server_response,
                        "timestamp": datetime.datetime.utcnow()
                    }

                    sessions_collection.update_one(
                        {"_id": ObjectId(session_id)},
                        {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
                    )
                    setter=0
                    return {
                        "user_message": user_message,
                        "message": server_response,
                        "file":server_response
                    }


            else:
                print("No process tag found.")

            
        
    else:
        import datetime
        print("clicked normal")
        global without

        global clause_list
        if without==0:
           
            without+=1

            text = user_message
            final_text=text

            contract_name = type_checker_agent(text)
            

            clause_list = clause_generation_agent(text, contract_name)
            print(clause_list) 
            clauses = clause_list
            context = text
            questions_dict = asker_agent1(clauses, context)
            global user_questions
            global web_questions
            user_questions = questions_dict.get("user", [])
            web_questions = questions_dict.get("web", [])
            print("User Questions:")
            print(user_questions)
            print("\n Web Questions:")
            print(web_questions)
            global web_answers
            web_answers = research_contract_questions(web_questions=web_questions)
            global q_len
            q_len = len(user_questions)
            global count_normal
            count_normal = 0 
            global ans_list
            ans_list =""
            
            user_message_entry = {
                "sender": "user",
                "text": user_message,
                "timestamp": datetime.datetime.utcnow()
            }

            server_response = user_questions[count_normal]  
            count_normal+=1

            server_message_entry = {
                "sender": "server",
                "text": server_response,
                "timestamp": datetime.datetime.utcnow()
            }

            sessions_collection.update_one(
                {"_id": ObjectId(session_id)},
                {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
            )
            setter=0
            return {
                "user_message": user_message,
                "message": server_response,
            }
        if without==1:
            if count_normal != q_len:
                input_text = user_message
                ans_list = ans_list + input_text + "\n"
                count_normal = count_normal + 1

                user_message_entry = {
                    "sender": "user",
                    "text": user_message,
                    "timestamp": datetime.datetime.utcnow()
                }


                try:
                    server_response = user_questions[count_normal]
                except IndexError:
                    print("hello")
                    without+=1
                    server_response = "All necessary information to draft the contract is now available. May I proceed to process the document?"
                    server_message_entry = {
                        "sender": "server",
                        "text": server_response,
                        "timestamp": datetime.datetime.utcnow()
                    }
                    sessions_collection.update_one(
                        {"_id": ObjectId(session_id)},
                        {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
                    )

                    return {
                        "user_message": user_message,
                        "message": server_response,
                    }
                    
                server_message_entry = {
                    "sender": "server",
                    "text": server_response,
                    "timestamp": datetime.datetime.utcnow()
                }

                sessions_collection.update_one(
                    {"_id": ObjectId(session_id)},
                    {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
                )

                return {
                    "user_message": user_message,
                    "message": server_response,
                }
            else:
                without+=1
                input_text = user_message
                user_message_entry = {
                    "sender": "user",
                    "text": user_message,
                    "timestamp": datetime.datetime.utcnow()
                }



                server_response = "All necessary information to draft the contract is now available. May I proceed to process the document?"  # Example: Simple echo response

                server_message_entry = {
                    "sender": "server",
                    "text": server_response,
                    "timestamp": datetime.datetime.utcnow()
                }

                sessions_collection.update_one(
                    {"_id": ObjectId(session_id)},
                    {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
                )

                return {
                    "user_message": user_message,
                    "message": server_response,
                }

        if without==2:
            Apercentage = 60 
            Bpercentage = 40
            pinecone_api_key="PINECONE API KEY"
            drafter = ContractDrafter(pinecone_api_key= pinecone_api_key, groq_api_key=apiaiconracter)
                
            contract = drafter.draft_full_contract(
                    contract_name=contract_name,
                    clause_order=clause_list,
                    requirements=final_text + ans_list + web_answers,
                    party_a_favorability=Apercentage,
                    party_b_favorability=Bpercentage
                )
            print(contract)
            user_message_entry = {
                    "sender": "user",
                    "text": user_message,
                    "timestamp": datetime.datetime.utcnow()
            }

            with open(CONTRACT_FILE_PATH, "w", encoding="utf-8") as file:
                file.write(format_legal_text(contract))
            print("HTML file saved as contract.html")

            server_response =CONTRACT_FILE_PATH 
            without+=1

            server_message_entry = {
                "sender": "server",
                "text": server_response,
                "file":server_response,
                "timestamp": datetime.datetime.utcnow()
            }

            sessions_collection.update_one(
                {"_id": ObjectId(session_id)},
                {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
            )

            return {
                "user_message": user_message,
                "message": "Would you like me to check for any corrections in the provided contract? I'll make the necessary changes, let you know once it's done, and then share the details of what was corrected.",
                "file":server_response,
                
            }
        if without==3:
            print("entered into without 3")
            user_message_entry = {
                    "sender": "user",
                    "text": user_message,
                    "timestamp": datetime.datetime.utcnow()
                }
            server_response = "Before we begin the negotiation, please let me know if you have any corrections you'd like to make to the contract. If so, please specify the changes you want to make here. kindly fill the form provided in the top right"  # Example: Simple echo response

            without+=1
            server_message_entry = {
                "sender": "server",
                "text": server_response,
                "timestamp": datetime.datetime.utcnow()
            }

            sessions_collection.update_one(
                {"_id": ObjectId(session_id)},
                {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
            )

            return {
                "user_message": user_message,
                "message": server_response,
            }
        if without==4:
            print("entered into without 4")
            intent2 = user_message
            text = confirmation_agent1(intent2)

            process_match = re.search(r"<process>(.*?)</process>", text)
            correction_match = re.search(r"<correction>(.*?)$", text, re.DOTALL)

            process_value = process_match.group(1) if process_match else None
            correction_value = correction_match.group(1).strip() if correction_match else None

            print("Process:", process_value)
            print("Correction:", correction_value)

            without_4_data=await request.json()
            contract_content=without_4_data.get("contract")
            global form_data
            form_data=without_4_data.get("FullDetails")
            popup_details=form_data


            original_document=html2text.html2text(contract_content)
            text = original_document
            chunks = re.split(r'###\s+(?=\w+)', text.strip())
            chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
            correction = correction_value
            chunks_count = len(chunks) 
            counter = 0
            final_list = []
            while (chunks_count != counter):
                corrected_chunk = correction_agent1(chunks[counter], correction)
                final_list.append(corrected_chunk)
                counter += 1
            print(final_list)
            full_Content = "\n".join(final_list)
            format_content=format_legal_text(full_Content)
            with open(CONTRACT_FILE_PATH, "w", encoding="utf-8") as f:
                f.write(format_content)

            user_message_entry = {
                    "sender": "user",
                    "text": user_message,
                    "timestamp": datetime.datetime.utcnow()
                }
            server_response = f"Okay, I've made the corrections you mentioned. Would you like to start the negotiation now? If there are any other changes you'd like to make before we begin, please let me know."  # Example: Simple echo response

            without+=1
            server_message_entry = {
                "sender": "server",
                "text": server_response,
                "timestamp": datetime.datetime.utcnow()
            }


            sessions_collection.update_one(
                {"_id": ObjectId(session_id)},
                {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
            )

            return {
                "user_message": user_message,
                "message": server_response,
                "file":CONTRACT_FILE_PATH
            }
        if without==5:
            print("entered into without 5")
            intent2 = user_message
            text = confirmation_agent1(intent2)
            process_match = re.search(r"<process>(.*?)</process>", text)
            correction_match = re.search(r"<correction>(.*?)$", text, re.DOTALL)

            process_value = process_match.group(1) if process_match else None
            correction_value = correction_match.group(1).strip() if correction_match else None

            print("Process:", process_value)
            print("Correction:", correction_value)

            

            if process_value == "YES":
                without+=1
                send_negotiation_email(form_data,CONTRACT_FILE_PATH)
                user_message_entry = {
                        "sender": "user",
                        "text": user_message,
                        "timestamp": datetime.datetime.utcnow()
                    }
                server_response = "I've sent the contract to the counterparty. Please check your email and let me know if you receive it!"

                server_message_entry = {
                    "sender": "server",
                    "text": server_response,
                    "timestamp": datetime.datetime.utcnow()
                }

  
                sessions_collection.update_one(
                    {"_id": ObjectId(session_id)},
                    {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
                )

                return {
                    "user_message": user_message,
                    "message": server_response,
                    
                }
            elif process_value == "NO":
                without+=1
            else:
                chunks_count = len(chunks) 
                counter = 0
                final_list = []
                while (chunks_count != counter):
                    corrected_chunk = correction_agent1(chunks[counter], correction)
                    final_list.append(corrected_chunk)
                    counter += 1

                full_Content = "\n".join(final_list)
                format_content=format_legal_text(full_Content)
                with open(CONTRACT_FILE_PATH, "w", encoding="utf-8") as f:
                    f.write(format_content)

                subprocess.run(["wkhtmltopdf", CONTRACT_FILE_PATH, pdf__path], check=True)

                user_message_entry = {
                        "sender": "user",
                        "text": user_message,
                        "timestamp": datetime.datetime.utcnow()
                    }
                server_response = f"Okay, I've made the corrections you mentioned. Would you like to start the negotiation now? If there are any other changes you'd like to make before we begin, please let me know."  # Example: Simple echo response

                server_message_entry = {
                    "sender": "server",
                    "text": server_response,
                    "timestamp": datetime.datetime.utcnow()
                }

                sessions_collection.update_one(
                    {"_id": ObjectId(session_id)},
                    {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
                )

                return {
                    "user_message": user_message,
                    "message": server_response,
                    "file":CONTRACT_FILE_PATH
                }
        if without==6:
            

            print("entered into without 6")

            EMAIL =  form_data.get("sender_email", "")
            PASSWORD ="iqoa qbdx wzve dolo"
            TARGET_SENDER = form_data.get("receiver_email", "")

            body = ""
            from datetime import timezone
            start_time = datetime.datetime.now(timezone.utc)

            def connect_to_mailbox():
                mail = imaplib.IMAP4_SSL("imap.gmail.com")
                mail.login(EMAIL, PASSWORD)
                mail.select("inbox")
                return mail

            def extract_and_print(msg):
                subject, encoding = decode_header(msg["Subject"])[0]
                if isinstance(subject, bytes):
                    subject = subject.decode(encoding or "utf-8")
                print("\n New Email from", TARGET_SENDER)
                print(" Subject:", subject)

                
                if msg.is_multipart():
                    for part in msg.walk():
                        if part.get_content_type() == "text/plain" and "attachment" not in str(part.get("Content-Disposition")):
                            body = part.get_payload(decode=True).decode(errors="ignore")
                            print("Body:\n", body)
                            user_message_entry = {
                                    "sender": "user",
                                    "text": user_message,
                                    "timestamp": datetime.datetime.utcnow()
                                }
                            server_response = f"the counter party replied /n {body}"  

                            server_message_entry = {
                                "sender": "server",
                                "text": server_response,
                                "timestamp": datetime.datetime.utcnow()
                            }
                        
                            sessions_collection.update_one(
                                {"_id": ObjectId(session_id)},
                                {"$push": {"messages": {"$each": [user_message_entry, server_message_entry]}}}
                            )

                            return {
                                "user_message": user_message,
                                "message": server_response,
                                
                            }
                            
                            
                else:
                    body = msg.get_payload(decode=True).decode(errors="ignore")
                    print(" Body:\n", body)

            def watch_new_emails_only():
                print(f" Watching new mails from {TARGET_SENDER} (starting {start_time})...")
                
                while True:
                    try:
                        mail = connect_to_mailbox()
                        status, data = mail.search(None, f'FROM "{TARGET_SENDER}"')
                        if status == "OK":
                            for e_id in data[0].split():
                                res, msg_data = mail.fetch(e_id, "(RFC822)")
                                for part in msg_data:
                                    if isinstance(part, tuple):
                                        msg = email.message_from_bytes(part[1])
                                        date_tuple = email.utils.parsedate_tz(msg["Date"])
                                        if date_tuple:
                                            email_time = datetime.datetime.fromtimestamp(email.utils.mktime_tz(date_tuple), tz=timezone.utc)

                                            if email_time > start_time:
                                                extract_and_print(msg)
                                                mail.logout()
                                                return  
                        mail.logout()
                        time.sleep(10)  

                    except imaplib.IMAP4.abort as err:
                        print("IMAP connection dropped. Reconnecting...")
                        time.sleep(5)
                    except Exception as e:
                        print(f"Unexpected Error: {e}")
                        time.sleep(5)
            without=without+1
            watch_new_emails_only()

@app.get("/get_contract/", response_class=HTMLResponse)
async def get_contract():
    setter=0
    if os.path.exists(CONTRACT_FILE_PATH):
        with open(CONTRACT_FILE_PATH, "r", encoding="utf-8") as file:
            html_content = file.read()
        print(html_content)
        return HTMLResponse(content=html_content, status_code=200)
    
    raise HTTPException(status_code=404, detail="Contract file not found")


@app.get("/get_messages/{session_id}")
async def get_messages(session_id: str):
    session = sessions_collection.find_one({"_id": ObjectId(session_id)})
    if not session:
        return {"error": "Session not found!"}

    return {"messages": session.get("messages", [])}

@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()
    while True:
        data = await websocket.receive_text()
        messages_collection.insert_one({"message": data})
        await websocket.send_text(f"Received: {data}")

@app.get("/get_sessions/")
async def get_sessions(request:Request):
    user_id = get_userid_from_request(request)
    print(user_id)
    template=False
    sessions = sessions_collection.find({"user_id": user_id}, {"_id": 1, "created_at": 1, "messages": 1})

    session_list = []
    for session in sessions:
        first_user_message = None


        if "messages" in session and isinstance(session["messages"], list):
            for msg in session["messages"]:
                if msg.get("sender") == "user":
                    first_user_message = msg.get("text")
                    break  
        
        session_list.append({
    "session_id": str(session["_id"]),
    "created_at": session.get("created_at", "Unknown"), 
    "first_message": first_user_message or "No messages yet"
})


    return session_list


if __name__ == "__main__":

    uvicorn.run(app, host="127.0.0.1", port=8000)

